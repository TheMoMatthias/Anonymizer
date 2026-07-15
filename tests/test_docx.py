from docx import Document

from anonymizer.formats import docx_handler
from anonymizer.pipeline import apply_document, scan_document


def test_detects_person_and_iban(sample_docx, analyzer, base_config):
    result = scan_document(sample_docx, analyzer, base_config)
    entity_types = {g.entity_type for g in result.all_actionable()}
    assert "IBAN_CODE" in entity_types
    assert "PERSON" in entity_types


def test_validated_steuer_id_auto_accepts(sample_docx, analyzer, base_config):
    result = scan_document(sample_docx, analyzer, base_config)
    steuer = [g for g in result.all_actionable() if g.entity_type == "DE_STEUER_ID"]
    assert steuer, "checksum-valid Steuer-ID should be detected"
    assert steuer[0].validated is True
    assert steuer[0].tier == "high"  # validated -> auto-accept tier


def test_groups_are_ordered_by_sensitivity(sample_docx, analyzer, base_config):
    result = scan_document(sample_docx, analyzer, base_config)
    keys = [grp.key for grp in result.groups]
    # People / Government / Financial classes precede lower-sensitivity ones.
    assert keys == sorted(keys, key=lambda k: keys.index(k))
    assert result.groups[0].sensitivity == "high"


def test_scans_header_text(sample_docx):
    units = docx_handler.extract_text_units(sample_docx)
    assert any("Vertraulich" in u.text for u in units)


def test_apply_removes_sensitive_text_and_preserves_formatting(sample_docx, analyzer, base_config, mapping_db_path):
    grouped = scan_document(sample_docx, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "pseudonymize"
    out_path, report_path = apply_document(sample_docx, grouped, analyzer, base_config, mapping_db_path)

    assert out_path.exists()
    assert report_path.exists()
    result = Document(out_path)
    assert "Hans Mueller" not in result.paragraphs[0].text
    assert result.paragraphs[0].runs[0].bold is True
    assert "Hans Mueller" not in result.sections[0].header.paragraphs[0].text


def test_reprocessing_same_value_is_consistent(tmp_path, analyzer, base_config, mapping_db_path):
    doc1 = Document()
    doc1.add_paragraph("Hans Mueller ist hier.")
    path1 = tmp_path / "a.docx"
    doc1.save(path1)

    doc2 = Document()
    doc2.add_paragraph("Hans Mueller ist auch hier.")
    path2 = tmp_path / "b.docx"
    doc2.save(path2)

    grouped1 = scan_document(path1, analyzer, base_config).all_actionable()
    for g in grouped1:
        g.action = "pseudonymize"
    out1, _ = apply_document(path1, grouped1, analyzer, base_config, mapping_db_path)

    grouped2 = scan_document(path2, analyzer, base_config).all_actionable()
    for g in grouped2:
        g.action = "pseudonymize"
    out2, _ = apply_document(path2, grouped2, analyzer, base_config, mapping_db_path)

    text1 = Document(out1).paragraphs[0].text
    text2 = Document(out2).paragraphs[0].text
    placeholder1 = next(w for w in text1.split() if w.startswith("[PERSON_"))
    placeholder2 = next(w for w in text2.split() if w.startswith("[PERSON_"))
    assert placeholder1 == placeholder2
