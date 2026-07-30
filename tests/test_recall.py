"""Recall coverage for the German-banking recognizers added in the Phase-4
audit: BIC/SWIFT, dates/DOB, postal addresses, bare customer numbers, and the
country-code-gated BIC validator used by the completeness backstop.

All scans force a single German language (`languages: ["de"]`) to mirror the
real per-document routing the pipeline applies.
"""

import pytest

from anonymizer import taxonomy
from anonymizer.core import build_scan_result, detect_unit
from anonymizer.models import TextUnit
from anonymizer.validators import bic_valid


def _findings(analyzer, base_config, text):
    cfg = {**base_config, "languages": ["de"]}
    return detect_unit(analyzer, TextUnit("u1", text), cfg)


def _types(analyzer, base_config, text):
    return {(f.entity_type, f.value) for f in _findings(analyzer, base_config, text)}


def _art9(analyzer, base_config, text):
    """Findings that land in the GDPR Art. 9 special-category data class."""
    return [
        f
        for f in _findings(analyzer, base_config, text)
        if taxonomy.data_class_for(f.entity_type).key == taxonomy.SPECIAL_CATEGORY.key
    ]


def test_bic_does_not_leak(analyzer, base_config):
    """A labeled BIC must be caught. It may classify as BIC_CODE or, when spaCy
    also tags the token as ORGANIZATION (which can outrank the context-gated
    BIC), under that category -- either way it is redacted, which is the property
    that matters. (Unlabeled BICs are covered by the completeness backstop.)"""
    typed = _types(analyzer, base_config, "Bitte überweisen an BIC: COBADEFFXXX zeitnah.")
    assert any("COBADEFF" in v for _et, v in typed)


def test_plain_uppercase_word_not_flagged_as_bic(analyzer, base_config):
    # No bic/swift context -> below threshold; must not flag an 8-letter word.
    typed = _types(analyzer, base_config, "Bitte das DOKUMENT prüfen und danach ABSENDEN.")
    assert not any(et == "BIC_CODE" for et, _v in typed)


def test_german_date_detected(analyzer, base_config):
    typed = _types(analyzer, base_config, "Geburtsdatum: 15.03.1980 des Kunden.")
    assert any(et == "DATE_TIME" and "15.03.1980" in v for et, v in typed)


def test_address_street_detected(analyzer, base_config):
    typed = _types(analyzer, base_config, "Anschrift: Musterstraße 12a in der Akte.")
    assert any(et == "DE_ADDRESS" and "Musterstraße" in v for et, v in typed)


def test_plz_city_detected(analyzer, base_config):
    typed = _types(analyzer, base_config, "Wohnort ist 50667 Köln laut Unterlagen.")
    assert any(et == "DE_ADDRESS" and "50667" in v for et, v in typed)


def test_kundennummer_detected_with_context(analyzer, base_config):
    typed = _types(analyzer, base_config, "Die Kundennummer 4830123 ist im System.")
    assert any(et == "DE_KUNDENNUMMER" and "4830123" in v for et, v in typed)


def test_anchored_name_detected_without_sentence_context(analyzer, base_config):
    """de_core_news_lg is WikiNER-trained, so it misses a name with no sentence
    context -- including the most common line in a German bank letter. The
    structural anchors must catch these, and must NOT swallow the honorific."""
    for text in ("Sehr geehrter Herr Müller,", "Name: Müller", "Kunde: Müller"):
        typed = _types(analyzer, base_config, text)
        assert ("PERSON", "Müller") in typed, f"missed the name in {text!r}: {typed}"


def test_honorific_is_not_part_of_the_name(analyzer, base_config):
    """'Herr Müller' and a bare 'Müller' must be ONE person, not two tokens."""
    typed = _types(analyzer, base_config, "Herr Müller hat das Konto eröffnet.")
    assert ("PERSON", "Müller") in typed


def test_misc_entities_surface_instead_of_being_dropped(analyzer, base_config):
    """Regression: spaCy tags 'Frau Bauer' as MISC; Presidio's mapping had no
    MISC key so the span was silently DISCARDED and the name leaked."""
    typed = _types(analyzer, base_config, "Frau Bauer zahlt.")
    assert any("Bauer" in v for _et, v in typed), f"MISC entity dropped: {typed}"


def test_lowercase_word_never_matches_case_sensitive_pattern(analyzer, base_config):
    """Regression: Presidio defaults to IGNORECASE, so the [A-Z] BIC pattern
    matched ordinary German words. At sensitivity 0.15 that redacted them."""
    cfg = {**base_config, "languages": ["de"], "sensitivity": 0.15}
    findings = detect_unit(analyzer, TextUnit("u1", "Sehr geehrter Herr, wie ausgefuehrt."), cfg)
    assert not any(f.entity_type == "BIC_CODE" for f in findings)


def test_name_propagates_across_the_document(tmp_path, analyzer, base_config):
    """The anchored salutation seeds the name; propagation must then catch the
    bare occurrences NER cannot see."""
    from docx import Document

    from anonymizer.pipeline import scan_document

    doc = Document()
    doc.add_paragraph("Sehr geehrter Herr Müller,")
    doc.add_paragraph("Die Unterlagen wurden von Müller geprüft.")
    doc.add_paragraph("Müller")
    path = tmp_path / "letter.docx"
    doc.save(path)

    persons = [g for g in scan_document(path, analyzer, base_config).all_actionable() if g.entity_type == "PERSON"]
    match = [g for g in persons if g.value == "Müller"]
    assert match, f"name not detected at all: {[(g.entity_type, g.value) for g in persons]}"
    assert match[0].count >= 3, f"propagation missed bare occurrences (count={match[0].count})"


def test_name_column_header_catches_bare_surnames(tmp_path, analyzer, base_config):
    """A spreadsheet column headed 'Name' is the one place a bare surname
    legitimately appears with no prose. NER only finds ~35% of ordinary German
    surnames there; the header is stronger evidence than the model."""
    import openpyxl

    from anonymizer.pipeline import scan_document

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Name"
    for row, surname in enumerate(["Müller", "Weber", "Bauer", "Koch"], start=2):
        ws[f"A{row}"] = surname
    path = tmp_path / "kunden.xlsx"
    wb.save(path)

    found = {g.value for g in scan_document(path, analyzer, base_config).all_actionable()}
    for surname in ("Müller", "Weber", "Bauer", "Koch"):
        assert surname in found, f"{surname} leaked from a 'Name' column: {found}"


def test_name_column_override_is_header_gated(tmp_path, analyzer, base_config):
    """The override must key off the HEADER, not blanket-flag every column. Uses
    a value spaCy ignores on its own ("Vorsorge"), so the header is the only
    variable: flagged under 'Name', untouched under 'Produktgruppe'."""
    import openpyxl

    from anonymizer.pipeline import scan_document

    def _scan(header: str) -> set[str]:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = header
        ws["A2"] = "Vorsorge"
        path = tmp_path / f"{header}.xlsx"
        wb.save(path)
        return {g.value for g in scan_document(path, analyzer, base_config).all_actionable()}

    assert "Vorsorge" not in _scan("Produktgruppe"), "override fired without a name header"
    assert "Vorsorge" in _scan("Name"), "override did not fire under a name header"


def test_bic_valid_country_gate():
    assert bic_valid("COBADEFFXXX")  # Commerzbank Frankfurt, DE
    assert bic_valid("DEUTDEFF")  # 8-char BIC, DE
    assert not bic_valid("TRANSFER")  # 'SF' not an ISO country
    assert not bic_valid("HELLO")  # wrong shape


# --- GDPR Art. 9 special-category data (German) -----------------------------
# A previous wave added the special_category CLASS to the taxonomy but no German
# recognizer ever emitted into it, so the class existed and never fired: health,
# religion and union/party data in a German workbook were a silent, total leak.


@pytest.mark.parametrize(
    "text,needle",
    [
        ("Diagnose: Diabetes mellitus", "Diabetes"),
        ("Krankenkasse: AOK Bayern", "AOK"),
        ("Pflegegrad: 3", "3"),
        ("GdB: 50", "50"),
        ("Der Kunde ist seit Mai arbeitsunfähig.", "arbeitsunfähig"),
        ("Schwerbehinderung liegt vor", "Schwerbehinderung"),
        ("Die Barmer bestätigt den Versicherungsschutz.", "Barmer"),
    ],
)
def test_art9_health_data_detected(analyzer, base_config, text, needle):
    """Health data is Art. 9 and had ZERO German detection."""
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), f"Art.9 health leak in {text!r}: {hits}"


@pytest.mark.parametrize(
    "text,needle",
    [
        ("Konfession: rk", "rk"),
        ("Kirchensteuermerkmal: ev", "ev"),
        ("Religionszugehörigkeit: islamisch", "islamisch"),
        ("Der Kunde ist evangelisch", "evangelisch"),
        ("Konfession: jüdisch", "jüdisch"),
    ],
)
def test_art9_religion_detected(analyzer, base_config, text, needle):
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), f"Art.9 religion leak in {text!r}: {hits}"


@pytest.mark.parametrize(
    "text,needle",
    [
        ("Gewerkschaft: ver.di", "ver.di"),
        ("Mitglied der IG Metall seit 2001", "IG Metall"),
        ("Parteimitgliedschaft: SPD", "SPD"),
        ("Gewerkschaftszugehörigkeit: DGB", "DGB"),
    ],
)
def test_art9_union_and_party_detected(analyzer, base_config, text, needle):
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), f"Art.9 union/party leak in {text!r}: {hits}"


@pytest.mark.parametrize(
    "text",
    [
        # A German bank workbook carries a Kirchensteuer amount column on EVERY
        # row. The AMOUNT is not the special-category datum and flagging it would
        # bury the reviewer -- whole-column policy is the right tool there.
        "Kirchensteuer 128,40 EUR",
        "Solidaritätszuschlag und Kirchensteuer werden einbehalten",
        # "die linke Seite" must not match the party "Die Linke" (case-sensitive).
        "Bitte die linke Spalte prüfen",
        # An ordinary payment sentence must not trip any Art. 9 pattern.
        "Die Abrechnung erfolgt quartalsweise gemäß den Geschäftsbedingungen.",
        # These probes run in the GERMAN scan (see _art9), where the German word
        # lists are active. English words quoted inside a German document must
        # still not collide with them: "aids" the verb and "orthodox" the
        # adjective are why HIV/AIDS is case-sensitive and why a bare "orthodox"
        # is not in the religion list. (An English-ROUTED document is covered
        # separately by test_german_art9_wordlists_do_not_fire_on_english.)
        "This tool aids the reconciliation process.",
        "We took an orthodox approach to the valuation.",
    ],
)
def test_art9_precision_does_not_flag_ordinary_bank_text(analyzer, base_config, text):
    hits = _art9(analyzer, base_config, text)
    assert not hits, f"Art.9 over-flagged {text!r}: {[(f.entity_type, f.value) for f in hits]}"


@pytest.mark.parametrize(
    "text,needle",
    [
        # spaCy claims the WHOLE phrase as one ORGANIZATION, i.e. a LONGER span that
        # CONTAINS the Art. 9 hit. _resolve_overlaps sorted by span length before
        # score, so the Art. 9 finding was dropped as "fully contained" and the value
        # was filed under Organizations & places -- whose action is PSEUDONYMIZE, a
        # reversible, mapping-backed [ORG_n]. That is verbatim the leak Art. 9
        # detection exists to close.
        ("Krankenkasse Barmer", "Barmer"),
        ("Gewerkschaft ver.di", "ver.di"),
        ("Versichert bei der AOK Bayern", "AOK"),
    ],
)
def test_art9_is_not_swallowed_by_a_longer_generic_span(analyzer, base_config, text, needle):
    """LEAK: special-category data must not lose its class to a longer generic
    ORGANIZATION/LOCATION span that merely contains it."""
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), (
        f"Art.9 value swallowed by a generic span in {text!r}: "
        f"{[(f.entity_type, f.value) for f in _findings(analyzer, base_config, text)]}"
    )
    # ...and the surviving set must still be non-overlapping, or apply splices the
    # text twice and produces a garbled token.
    spans = sorted((f.start, f.end) for f in _findings(analyzer, base_config, text))
    assert all(a[1] <= b[0] for a, b in zip(spans, spans[1:])), f"overlapping spans: {spans}"


def test_art9_stays_in_the_review_tier(analyzer, base_config):
    """Art. 9 detection is contextual, not checksummed, so it must land BELOW the
    auto-accept bar: over-flagging religion/health is as damaging as missing it,
    and the reviewer is the only thing that can tell them apart."""
    high = float(base_config.get("tiers", {}).get("high", 0.9))
    for text in ("Diagnose: Diabetes mellitus", "Konfession: rk", "Gewerkschaft: ver.di"):
        for f in _art9(analyzer, base_config, text):
            assert f.score < high, f"Art.9 finding auto-accepted: {f.entity_type} {f.value} {f.score}"


# --- C2: standalone (context-free) Sozialversicherungsnummer -----------------


def test_sv_nummer_detected_in_a_bare_cell(analyzer, base_config):
    """The SV recognizer was context-gated, so a bare SV number in its own
    spreadsheet column -- exactly how an HR workbook stores it -- was missed
    entirely. The check digit makes a standalone match safe."""
    typed = _types(analyzer, base_config, "65170839J003")
    assert any(et == "DE_SV_NUMMER" for et, _v in typed), f"bare SV-Nummer leaked: {typed}"


def test_sv_nummer_bare_is_auto_accept_tier(analyzer, base_config):
    """A checksum-validated SV number is near-certainly real -> high tier."""
    hits = [f for f in _findings(analyzer, base_config, "65170839J003") if f.entity_type == "DE_SV_NUMMER"]
    assert hits and hits[0].validated is True and hits[0].score >= 0.9


def test_sv_nummer_failed_checksum_still_surfaces(analyzer, base_config):
    """A checksum-failing SV number is usually a TYPO in a real one, not a false
    positive -- it must stay visible for review, not be silently dropped."""
    hits = [f for f in _findings(analyzer, base_config, "65170839J004") if f.entity_type == "DE_SV_NUMMER"]
    assert hits, "checksum-failed SV-Nummer dropped instead of demoted"
    assert hits[0].validated is False


# --- C3: a checksum-FAILED structured ID must keep its data class ------------


def _classes_for(analyzer, base_config, text):
    cfg = {**base_config, "languages": ["de"]}
    unit = TextUnit("u1", text)
    findings = detect_unit(analyzer, unit, cfg)
    return findings, {g.key for g in build_scan_result(findings, [unit], cfg).groups}


@pytest.mark.parametrize(
    "entity_type,value",
    [
        # The IBAN case was xfail: spaCy claims the IDENTICAL span as NER_MISC at
        # its flat 0.85 and beat the 0.4-demoted IBAN_CODE on score. core's overlap
        # resolution now prefers the CHECKSUM-TESTED recognizer on an equal-length
        # span, so this passes -- the marker is gone rather than left to xpass
        # silently (see test_core.test_resolve_overlaps_checksum_tested_id_beats_...).
        ("IBAN_CODE", "DE89370400440532013001"),  # mod-97 fails on the last digit
        ("CREDIT_CARD", "4111 1111 1111 1112"),  # Luhn fails on the last digit
    ],
)
def test_checksum_failed_id_keeps_its_financial_data_class(analyzer, base_config, entity_type, value):
    """LEAK: Presidio's own recognizers score a failed checksum to 0 and then
    DROP the result (`if score > MIN_SCORE`), so a typo'd/OCR'd IBAN or card
    never reached `_refine`. It lost its Financial-IDs class entirely -- landing
    either in "Other named entities" (if spaCy happened to tag it) or in the
    informational possible-miss bucket the review UI never applies. A failed
    checksum is a typo in a REAL id, not a false positive: it must stay
    classified and reviewable."""
    findings, classes = _classes_for(analyzer, base_config, value)
    hits = [f for f in findings if f.entity_type == entity_type]
    assert hits, f"checksum-failed {entity_type} dropped: {[(f.entity_type, f.value) for f in findings]}"
    assert hits[0].validated is False
    assert taxonomy.FINANCIAL_IDS.key in classes, f"lost its data class, groups={classes}"


# --- C4: German compound / abbreviated / prepositional street names ----------


@pytest.mark.parametrize(
    "text,needle",
    [
        ("Rudolf-Breitscheid-Str. 12", "Rudolf-Breitscheid-Str."),
        ("Zum Alten Forsthaus 4a", "Zum Alten Forsthaus"),
        ("Am Wall 3", "Am Wall"),
        ("Kirchstr.7", "Kirchstr."),
        # The whole range must be covered -- "Hauptstraße 12" left "14" in place.
        ("Hauptstraße 12-14", "12-14"),
        ("Unter den Linden 5", "Unter den Linden"),
        ("50667 Köln-Ehrenfeld", "Köln-Ehrenfeld"),
        ("D-50667 Köln", "50667 Köln"),
    ],
)
def test_german_address_variants_detected(analyzer, base_config, text, needle):
    typed = _types(analyzer, base_config, text)
    assert any(et == "DE_ADDRESS" and needle in v for et, v in typed), f"address leak in {text!r}: {typed}"


@pytest.mark.parametrize(
    "text",
    [
        "Im Jahr 2024 wurden die Gebühren angepasst",
        "Im Anhang 2024 finden Sie die Aufstellung",
    ],
)
def test_prepositional_address_does_not_match_prose(analyzer, base_config, text):
    typed = _types(analyzer, base_config, text)
    assert not any(et == "DE_ADDRESS" for et, _v in typed), f"address over-flagged {text!r}: {typed}"


# --- C5: date-of-birth formats ----------------------------------------------


@pytest.mark.parametrize(
    "text,needle",
    [
        ("geb. 31.12.80", "31.12.80"),
        ("Geburtsdatum: 1980-12-31", "1980-12-31"),
        ("Müller, *1980", "1980"),
        ("*31.12.1980", "31.12.1980"),
        ("Jahrgang 1980", "1980"),
        ("31.12.1980", "31.12.1980"),
        ("geboren am 3. Mai 1980", "3. Mai 1980"),
    ],
)
def test_date_of_birth_variants_detected(analyzer, base_config, text, needle):
    """A DOB is a strong quasi-identifier; the German model emits no DATE at all,
    so anything the DATE_TIME pattern misses leaks silently."""
    typed = _types(analyzer, base_config, text)
    assert any(et == "DATE_TIME" and needle in v for et, v in typed), f"DOB leak in {text!r}: {typed}"


# --- C6: the recall harness must measure the BARE cell bare ------------------


# --- R2/P3: a LABEL:VALUE Art. 9 hit must claim to a BOUNDARY -----------------


@pytest.mark.parametrize(
    "text,label",
    [
        # A free-text Diagnose/Bemerkung cell: the old fixed 40-char cap cut
        # MID-TOKEN and left the tail of the sensitive note in the document
        # ("[X]nz mit begleitender [X] und Schlafstoerung").
        (
            "Diagnose: Verdacht auf chronische Niereninsuffizienz mit begleitender "
            "Depression und Schlafstoerung",
            "Diagnose: ",
        ),
        ("Krankenkasse: AOK Bayern zahlt den Beitrag ab Januar 2024 nicht mehr", "Krankenkasse: "),
        ("Konfession: evangelisch-lutherisch seit der Taufe im Mai neunzehnhundert", "Konfession: "),
        ("Gewerkschaft: Vereinte Dienstleistungsgewerkschaft Bezirk Nordrhein seit 2004", "Gewerkschaft: "),
    ],
)
def test_art9_label_value_claims_to_a_boundary_not_a_char_count(analyzer, base_config, text, label):
    """FALSELY-CLEAN OUTPUT: the label:value patterns capped the claimed value at
    40 (health) / 30 (religion) characters and cut mid-word, so a long free-text
    cell was only PARTIALLY redacted and the tail of the sensitive note survived
    into the "anonymized" file. The value must run to a real boundary -- end of
    cell / sentence / line -- never to a character count."""
    hits = _art9(analyzer, base_config, text)
    assert hits, f"no Art.9 finding at all in {text!r}"
    assert any(f.start == len(label) and f.end == len(text) for f in hits), (
        f"value truncated -- residual would remain: {[(f.value, f.start, f.end) for f in hits]}"
    )


@pytest.mark.parametrize(
    "text",
    [
        # A negation or a placeholder is not special-category data; claiming it
        # produced meaningless one-way-anonymized findings and pure reviewer noise.
        "Partei: keine",
        "Konfession: -",
        "Diagnose: k.A.",
        "Gewerkschaft: nein",
        "Krankenkasse: unbekannt",
        "Religion: keine Angabe",
    ],
)
def test_art9_label_value_ignores_negations_and_placeholders(analyzer, base_config, text):
    hits = _art9(analyzer, base_config, text)
    assert not hits, f"placeholder claimed as Art.9 data in {text!r}: {[(f.entity_type, f.value) for f in hits]}"


# --- R2/P4: German exports transliterate umlauts (ae/oe/ue) -------------------


@pytest.mark.parametrize(
    "text,needle",
    [
        ("arbeitsunfaehig", "arbeitsunfaehig"),
        ("Arbeitsunfaehigkeit seit Mai", "Arbeitsunfaehigkeit"),
        ("erwerbsunfaehig", "erwerbsunfaehig"),
        ("berufsunfaehig", "berufsunfaehig"),
        ("Der Kunde ist juedisch", "juedisch"),
        ("Mitglied bei Die Gruenen", "Gruenen"),
        ("Religionszugehoerigkeit: islamisch", "islamisch"),
    ],
)
def test_art9_transliterated_umlauts_are_detected(analyzer, base_config, text, needle):
    """The word lists were UMLAUT-ONLY, so the ae/oe/ue spelling that German
    system exports routinely use (this repo's own fixtures say "Mueller") was
    missed entirely -- a pure transliteration gap, not a vocabulary one."""
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), f"transliterated Art.9 term missed in {text!r}: {hits}"


# --- R2/P5: Art. 9 recall gaps (no-colon labels, and 3 missing categories) ----


@pytest.mark.parametrize(
    "text,needle",
    [
        # Word/PDF prose carries no ':' -- only the Excel handler synthesises one.
        ("Pflegegrad 3 wurde bewilligt", "3"),
        ("GdB 50 festgestellt", "50"),
        ("Krankenkasse TK", "TK"),
        ("Reha beantragt", "Reha"),
        ("Kirchensteuermerkmal ev", "ev"),
    ],
)
def test_art9_label_without_a_colon_is_detected(analyzer, base_config, text, needle):
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), f"Art.9 leak in colon-free prose {text!r}: {hits}"


@pytest.mark.parametrize(
    "text,needle",
    [
        # Racial/ethnic origin, sex life/orientation and genetic/biometric data are
        # Art. 9 categories that had NO detection at all.
        ("Herkunft: tuerkisch", "tuerkisch"),
        ("Staatsangehoerigkeit: syrisch", "syrisch"),
        ("Staatsangehörigkeit: tunesisch", "tunesisch"),
        ("Der Kunde ist homosexuell", "homosexuell"),
        ("Sexuelle Orientierung: bisexuell", "bisexuell"),
        ("Gentest liegt vor", "Gentest"),
        ("Fingerabdruck hinterlegt", "Fingerabdruck"),
    ],
)
def test_art9_origin_orientation_and_genetic_data_detected(analyzer, base_config, text, needle):
    hits = _art9(analyzer, base_config, text)
    assert any(needle in f.value for f in hits), f"Art.9 category with no detection at all: {text!r} -> {hits}"


# --- R2/P9-en: German word lists must not fire on an ENGLISH document ---------


@pytest.mark.parametrize(
    "text",
    [
        "The Great Depression started in 1929.",
        "Diabetes research fund performance",
        "The AOK index of the region",
    ],
)
def test_german_art9_wordlists_do_not_fire_on_english(analyzer, base_config, text):
    """PRECISION: Art. 9 findings are ONE-WAY anonymized (irreversible). German
    word lists firing on ordinary English financial vocabulary destroys text in
    an English document with no way back."""
    cfg = {**base_config, "languages": ["en"]}
    hits = [
        f
        for f in detect_unit(analyzer, TextUnit("u1", text), cfg)
        if f.entity_type.startswith("DE_")
    ]
    assert not hits, f"German recognizer fired on English text {text!r}: {[(f.entity_type, f.value) for f in hits]}"


# --- R2/P6: the prepositional address pattern vs German bank boilerplate ------


@pytest.mark.parametrize(
    "text",
    [
        "Zum Stichtag 31.12.2024 betrug der Saldo",
        "In der Anlage 3 finden Sie die Aufstellung",
        "Zum Beispiel 3 Positionen wurden geprueft",
        "Am Ende 3 Positionen",
    ],
)
def test_prepositional_address_does_not_match_bank_boilerplate(analyzer, base_config, text):
    """'Zum Stichtag' appears in essentially every German bank statement. A
    hand-maintained prose-exclusion list was not converging, so the pattern needs
    a positive signal: an address occupies a delimited segment, prose does not."""
    typed = _types(analyzer, base_config, text)
    assert not any(et == "DE_ADDRESS" for et, _v in typed), f"address over-flagged {text!r}: {typed}"


def test_stichtag_date_is_not_swallowed_by_a_bogus_address(analyzer, base_config):
    """The bogus 'Zum Stichtag 31.12.2024' address CROSSED the DATE_TIME finding
    and swallowed it in overlap resolution, so the date lost its own class."""
    typed = _types(analyzer, base_config, "Zum Stichtag 31.12.2024 betrug der Saldo")
    assert ("DATE_TIME", "31.12.2024") in typed, typed


# --- R2/P8: the CREDIT_CARD fallback must not claim internal reference runs ---


def test_bare_16_digit_reference_is_not_claimed_as_a_card(analyzer, base_config):
    """Roughly 15-20% of all 16-digit numbers carry a card-shaped prefix. As an
    unconditional fallback that was an ACTIONABLE finding which no configuration
    could filter (validated=False bypasses the confidence threshold), with a
    ONE-WAY default action -- so bulk-accepting Financial IDs destroyed internal
    16-digit reference numbers with no way back."""
    typed = _types(analyzer, base_config, "Vorgangsnummer 4123456789012345 im System")
    assert not any(et == "CREDIT_CARD" for et, _v in typed), f"internal reference claimed as a card: {typed}"


@pytest.mark.parametrize(
    "text",
    [
        "Kreditkarte 4123456789012345 wurde belastet",  # Luhn-INVALID, label-anchored
        "4123 4567 8901 2345",  # Luhn-INVALID, written in card groups
    ],
)
def test_checksum_failed_card_is_still_caught_when_it_looks_like_a_card(analyzer, base_config, text):
    typed = _types(analyzer, base_config, text)
    assert any(et == "CREDIT_CARD" for et, _v in typed), f"card number leaked: {typed}"


# --- R2/P10: report typography must not read as a date of birth --------------


@pytest.mark.parametrize("text", ["* 2024 Prognose", "Umsatz * 2019 nach Segment"])
def test_footnote_asterisk_is_not_a_birth_date(analyzer, base_config, text):
    """The genealogical birth marker is written tight ("*1980"); a spaced '*' is a
    footnote marker, which appears all over a report's typography."""
    typed = _types(analyzer, base_config, text)
    assert not any(et == "DATE_TIME" for et, _v in typed), f"footnote marker read as a DOB: {typed}"


# --- R2/P1: the non-validating IBAN fallback must not eat the next word ------


@pytest.mark.parametrize(
    "text,iban",
    [
        # Every country whose IBAN length is a multiple of 4 exposed the bug: the
        # fallback's optional trailing group could start with a SPACE, so it ate the
        # first character(s) of the next word.
        ("IBAN AT611904300234573201 Betrag 100 EUR", "AT611904300234573201"),  # AT = 20
        ("Konto BE68539007547034 Ueberweisung", "BE68539007547034"),  # BE = 16
        ("IBAN SE4550000000058398257466 Empfaenger", "SE4550000000058398257466"),  # SE = 24
        # A NUMBER following the IBAN is the case that matters most and the one an
        # earlier repair missed: narrowing the trailing group to digits while LEAVING
        # the optional space kept the defect alive in its worst form, because what
        # follows an IBAN on a German transfer line is the AMOUNT, not a word. The
        # tail then ate the amount and applying the finding DELETED it from the
        # document. Keep both shapes parameterised -- testing only the word case is
        # what let the incomplete fix through.
        ("IBAN AT611904300234573201 100 EUR", "AT611904300234573201"),
        ("Konto BE68539007547034 100 EUR", "BE68539007547034"),
        ("Betrag zu IBAN LU280019400644750000 25 Euro", "LU280019400644750000"),  # LU = 20
        ("IBAN AT611904300234573201 12 EUR", "AT611904300234573201"),
        # A German IBAN (22) has a 2-char remainder, so the tail is genuinely needed.
        ("IBAN DE89370400440532013000 ist das Konto", "DE89370400440532013000"),
        ("IBAN DE89 3704 0044 0532 0130 00 ist das Konto", "DE89 3704 0044 0532 0130 00"),
    ],
)
def test_iban_span_never_bleeds_into_the_following_word(analyzer, base_config, text, iban):
    """CORRUPTION + demotion regression: the fallback IBAN pattern claimed
    '<IBAN> B' of 'AT61... Betrag'. That longer bogus span beat presidio's correct
    one in overlap resolution, iban_valid() then failed on the polluted string and
    a CHECKSUM-VALID IBAN was demoted to 0.4/validated=False -- dropping it out of
    the auto-accept tier AND splicing one character out of the following word when
    the redaction was applied ('Betrag' -> 'etrag')."""
    hits = [f for f in _findings(analyzer, base_config, text) if f.entity_type == "IBAN_CODE"]
    assert hits, f"IBAN not detected at all in {text!r}"
    assert hits[0].value == iban, f"span bled into the next word: {hits[0].value!r}"
    assert hits[0].validated is True, "a checksum-VALID IBAN was demoted"
    assert hits[0].score >= 0.9, f"valid IBAN fell out of the auto-accept tier: {hits[0].score}"
    end = text.index(iban) + len(iban)
    assert all(f.end <= end or f.start >= end for f in _findings(analyzer, base_config, text)), (
        "a finding straddles the end of the IBAN -- applying it would corrupt the next word"
    )


@pytest.mark.parametrize(
    "text",
    [
        "ISIN DE0007164600 Stueckzahl 100",
        "Wertpapierkennnummer DE000BAY0017 Bestand",
        # A length FLOOR alone does not close this: with a space-absorbing trailing
        # group the 12-char ISIN simply reached the floor by swallowing the position
        # QUANTITY that follows it -- and 'ISIN <id> <quantity>' is the standard row
        # of a depot statement, so this was the common case, not an edge case.
        # Applying such a finding deleted the quantity from the statement.
        "ISIN DE0007164600 100 Stueck",
        "ISIN DE0007164600 1000 Stueck",
        "DE0007164600 250",
        "Depot: DE000BAY0017 500 Stueck",
    ],
)
def test_securities_identifier_is_not_claimed_as_an_iban(analyzer, base_config, text):
    """A 12-character ISIN is IBAN-SHAPED but can never be an IBAN (the shortest
    IBAN is 15). Claiming it flagged every securities position in a depot
    statement as a Financial ID."""
    typed = _types(analyzer, base_config, text)
    assert not any(et == "IBAN_CODE" for et, _v in typed), f"ISIN claimed as an IBAN: {typed}"


def test_same_iban_with_and_without_a_trailing_amount_is_one_finding(analyzer, base_config):
    """Referential-consistency regression. While the fallback span could absorb a
    following number, ONE account number produced TWO findings -- the polluted
    'AT61...3201 100' on a transfer line and the clean 'AT61...3201' elsewhere --
    so the same account was pseudonymized to [IBAN_1] in one place and [IBAN_2] in
    another, breaking precisely the cross-document consistency the mapping exists
    to provide."""
    text = (
        "Ueberweisung an IBAN AT611904300234573201 100 EUR.\n"
        "Das Konto AT611904300234573201 wurde geschlossen."
    )
    ibans = {f.value for f in _findings(analyzer, base_config, text) if f.entity_type == "IBAN_CODE"}
    assert ibans == {"AT611904300234573201"}, f"one account number split into {ibans}"


def test_bare_cell_probe_carries_no_filler_context():
    """The harness prepended neutral filler prose to EVERY probe, so the
    "bare_cell" stratum -- a lone value in a spreadsheet cell, the hardest and
    most common shape in the user's real workbooks -- was never measured bare.
    Every bare_cell number previously reported was therefore optimistic."""
    from anonymizer import evaluation

    assert evaluation.probe_text("bare_cell", "Anna", "Müller") == "Müller"
    assert evaluation.FILLER in evaluation.probe_text("prose_oblique", "Anna", "Müller")


# --- 2026-07-27 audit of a real internal workbook (docs/run_precision-rework_
# 2026-07-27.md). Every case below is a value that LEAKED out of that file. ---


def _wb_scan(tmp_path, analyzer, base_config, cells, name="audit.xlsx"):
    """Builds a one-sheet workbook from {coord: value} and returns the actionable
    finding values, so a leak is asserted end-to-end through the real pipeline
    rather than against a single recognizer."""
    import openpyxl

    from anonymizer.pipeline import scan_document

    wb = openpyxl.Workbook()
    ws = wb.active
    for coord, value in cells.items():
        ws[coord] = value
    path = tmp_path / name
    wb.save(path)
    return {g.value for g in scan_document(path, analyzer, base_config).all_actionable()}


def test_english_people_column_headers_catch_names(tmp_path, analyzer, base_config):
    """The measured recall gap: _NAME_HEADER_TERMS was German-only, so every name
    under Owner / Einreicher / MDX_Lead / MDX_Proxy leaked. 83 distinct people in
    221 cells went out in the clear on the reported workbook -- including full
    names, which the German NER model also missed in a bare cell."""
    found = _wb_scan(tmp_path, analyzer, base_config, {
        "A1": "Owner", "A2": "Ulf Gericke",
        "B1": "Einreicher", "B2": "Cordula",
        "C1": "MDX_Lead", "C2": "Marco",
        "D1": "MDX_Proxy", "D2": "Mirijam",
    })
    for name in ("Ulf Gericke", "Cordula", "Marco", "Mirijam"):
        assert name in found, f"{name!r} leaked from its own people column: {found}"


def test_url_is_detected(tmp_path, analyzer, base_config):
    """41 internal links leaked from the reported workbook. Presidio's
    UrlRecognizer was always loaded -- URL was simply not in the `entities` block,
    and detect_unit only requests configured entities, so it was never asked for."""
    found = _wb_scan(tmp_path, analyzer, base_config, {
        "A1": "Confluence Link",
        "A2": "https://emma.intern.example.com/x/Oim3JQ",
    })
    assert "https://emma.intern.example.com/x/Oim3JQ" in found, found


def test_url_entity_is_configured_and_labelled():
    """A URL finding must reach a data class and a token label, or it renders as
    an unreversible/unlabelled token."""
    import yaml

    from anonymizer.actions import TOKEN_LABELS
    from anonymizer.config import DEFAULT_CONFIG_PATH

    shipped = yaml.safe_load(DEFAULT_CONFIG_PATH.read_text(encoding="utf-8"))
    assert "URL" in shipped["entities"]
    assert TOKEN_LABELS["URL"] == "LINK"
    assert taxonomy.data_class_for("URL").key == taxonomy.BANK_INTERNAL.key


def test_name_fused_to_an_excel_escape_is_still_found(tmp_path, analyzer, base_config):
    """Excel writes a control character as a literal `_xHHHH_` escape and openpyxl
    passes it through verbatim, so a multi-value cell reads as
    "Ulf Gericke_x001E_". Before neutralization the fused token was rejected
    outright by _is_structural_nonname's underscore rule -- a silent false
    negative on a name column."""
    found = _wb_scan(tmp_path, analyzer, base_config, {
        "A1": "Owner", "A2": "Ulf Gericke_x001E_",
    })
    assert "Ulf Gericke" in found, f"the escape hid the name: {found}"


def test_a_weaker_guess_cannot_veto_a_people_column_header(tmp_path, analyzer, base_config):
    """spaCy types some real names NER_MISC rather than PERSON. That whole-cell
    MISC hit suppressed the header override, and MISC -- a bare guess -- was then
    dropped outright by corroboration_only, so the name left in the clear from a
    column literally headed "Owner". Measured on the reported workbook with
    'Constanza Hiemenz'."""
    from anonymizer.formats.xlsx_handler import _analyze_cell_text

    cfg = {**base_config, "languages": ["de"]}
    findings = _analyze_cell_text("Constanza Hiemenz", "Einreicher", analyzer, cfg)
    assert [(f.entity_type, f.source) for f in findings] == [("PERSON", "whole_cell_override")], findings

    # No people header -> no retype. The header is what carries the evidence, so
    # this must not become a blanket "MISC is really PERSON" rule.
    plain = _analyze_cell_text("Constanza Hiemenz", "Notiz", analyzer, cfg)
    assert [f.entity_type for f in plain] == ["NER_MISC"], plain

    found = _wb_scan(tmp_path, analyzer, base_config, {
        "A1": "Einreicher", "A2": "Constanza Hiemenz",
    }, name="misc_name.xlsx")
    assert "Constanza Hiemenz" in found, found


# --- Art.9 bare-value gaps found by scripts/measure_recall.py (2026-07-27) ----


@pytest.mark.parametrize("text,needle", [
    # A German HR sheet writes the diagnosis BARE in the cell; the label lives in the
    # column header, which the xlsx handler prepends but docx/pdf prose never does.
    ("neuapostolisch", "neuapostolisch"),
    ("Bandscheibenvorfall L4/L5", "Bandscheibenvorfall"),
    ("chronische Migraene", "Migraene"),
    ("Herzinfarkt im Mai", "Herzinfarkt"),
])
def test_art9_bare_values_that_were_leaking_are_detected(analyzer, base_config, text, needle):
    """These were measured as leaks in the BARE form -- the anchored 'Konfession: X'
    and 'Diagnose: X' shapes already worked. Art.9 is the most damaging class to
    miss, so a bare-cell gap here matters more than anywhere else."""
    found = {f.value for f in _art9(analyzer, base_config, text)}
    assert any(needle in v for v in found), f"{needle!r} leaked from {text!r}: {found}"


@pytest.mark.parametrize("text", [
    # The bare lists match UNINFLECTED forms only, which is what keeps organisation
    # names intact -- \b fails on the trailing inflection. This is the property that
    # makes the lists safe to extend at all, so it is pinned explicitly.
    "Neuapostolische Kirchengemeinde Frankfurt",
    "Evangelische Bank eG",
    "Die Katholische Universitaet Eichstaett",
    "Methodistische Gemeinde Berlin",
    # Deliberately absent from the word lists -- each would one-way destroy ordinary
    # text, which is unrecoverable.
    "Herr Krebs hat angerufen",
    "Covid-Massnahmen wurden aufgehoben",
])
def test_art9_bare_lists_do_not_claim_org_names_or_ordinary_prose(analyzer, base_config, text):
    assert not _art9(analyzer, base_config, text), text


def test_propagation_does_not_erase_an_anchor_it_wins_over(analyzer, base_config):
    """Propagation scores 0.85 while the anchored name patterns score 0.70-0.75, so a
    propagated occurrence WINS the overlap on the very span an anchor corroborated.
    _absorb_corroborating_source only transferred a dropped source when the KEPT one was
    spaCy, so that anchor was silently destroyed -- every occurrence ended up
    source="propagation", the group read as a bare guess, and under corroboration-only
    the name was demoted and LEAKED. Propagation must not create corroboration, but it
    must not erase it either."""
    from docx import Document

    from anonymizer.pipeline import scan_document

    import tempfile
    from pathlib import Path
    with tempfile.TemporaryDirectory() as td:
        doc = Document()
        doc.add_paragraph("Sehr geehrter Herr Müller,")
        doc.add_paragraph("Die Unterlagen wurden von Müller geprüft.")
        path = Path(td) / "letter.docx"
        doc.save(path)
        result = scan_document(path, analyzer, base_config)

    persons = [g for g in result.all_actionable() if g.entity_type == "PERSON"]
    assert any(g.value == "Müller" for g in persons), (
        f"the anchored salutation must keep the name actionable: "
        f"actionable={[(g.entity_type, g.value) for g in result.all_actionable()]} "
        f"demoted={[(g.entity_type, g.value) for g in result.demoted]}"
    )


def test_english_honorifics_are_stripped_and_the_name_corroborated(analyzer, base_config):
    """engine._HONORIFICS recognised Mr/Mrs/Ms all along, but both STRIPPERS were
    German-only -- so an English name kept its title inside the finding value. That
    keyed the pseudonym on the title AND stopped the given-name gazetteer (which tests
    the first token) from corroborating it, so under corroboration-only the name was
    demoted and leaked."""
    cfg = {**base_config, "languages": ["en"]}
    found = {f.value: f.source for f in _findings(analyzer, cfg, "Ms Priya Whitfield approved the request.")}
    assert "Priya Whitfield" in found, f"honorific not stripped: {found}"
    assert "Ms Priya Whitfield" not in found


def test_a_german_genitive_inherits_its_base_name_corroboration():
    """"Kochs Team" is the same person as "Koch", but NER reports it as its own value,
    so it formed its own uncorroborated group. Demoting it left the surname legible
    while "Koch" was redacted everywhere else -- a leak the fail-loud verify caught."""
    from anonymizer import core
    from anonymizer.models import Finding as F

    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            F("PERSON", "Koch", 0.85, "c", "u1", 0, 4, source="whole_cell_override"),
            F("PERSON", "Kochs", 0.85, "c", "u2", 0, 5, source="SpacyRecognizer"),
            # A word merely ENDING in s must not inherit from a stem that is not itself
            # a corroborated name.
            F("PERSON", "Prozess", 0.85, "c", "u3", 0, 7, source="SpacyRecognizer"),
        ],
        [TextUnit("u", "x")], cfg,
    )
    vals = {g.value for g in result.all_actionable()}
    # With PERSON now corroboration-only (2026-07-30) this test finally exercises the
    # rule it was written for: "Koch" is corroborated by the column override, "Kochs"
    # INHERITS that corroboration as its genitive, and "Prozess" -- a word that merely
    # ends in s, with no corroborated stem -- is correctly demoted.
    assert vals == {"Koch", "Kochs"}, vals
    assert {g.value for g in result.demoted} == {"Prozess"}, [g.value for g in result.demoted]


# --- 2026-07-30: the hardened recall harness and what it exposed --------------
# Each test below pins a leak the harness measured once the easy contexts stopped
# being the only ones tested. See docs/run_precision-rework_2026-07-27.md.


def test_harness_drops_particles_but_requires_both_halves_of_a_hyphenated_name():
    """The two asymmetric scoring rules the hard strata depend on. Getting either
    backwards makes the whole report lie: dropping a hyphen half would credit a
    genuine leak ("Rottluff" left standing), and requiring the particle would
    punish correct behaviour (a bare "von" discloses nobody)."""
    from anonymizer.evaluation import _identifying_tokens

    assert _identifying_tokens("von der Leyen") == ["Leyen"]
    assert _identifying_tokens("de la Cruz") == ["Cruz"]
    assert _identifying_tokens("Schmidt-Rottluff") == ["Schmidt", "Rottluff"]
    assert _identifying_tokens("Koch") == ["Koch"]


def test_harness_refuses_to_credit_a_half_caught_hyphenated_name():
    """A run that redacted every "Schmidt" and no "Rottluff" has protected nobody."""
    from anonymizer.evaluation import _found
    from anonymizer.models import Finding as F

    half = [F("PERSON", "Schmidt", 0.9, "c", "u1", 0, 7)]
    both = half + [F("PERSON", "Rottluff", 0.9, "c", "u2", 0, 8)]
    assert not _found(half, "Schmidt-Rottluff")
    assert _found(both, "Schmidt-Rottluff")


@pytest.mark.parametrize(
    "text,expected",
    [
        # The particle name was the single worst anchor failure: _NAME required
        # every token to start uppercase, so the most common line in a German
        # bank letter matched NOTHING for a customer with a nobiliary particle.
        ("Sehr geehrter Herr von Bergen,", "von Bergen"),
        ("Sehr geehrte Frau van den Broek,", "van den Broek"),
        ("Kunde: de la Cruz", "de la Cruz"),
        # Naming somebody by ROLE rather than title scored 0% for German
        # common-noun surnames before these anchors existed.
        ("Der Einreicher Winkler aus Frankfurt bestätigte den Sachverhalt.", "Winkler"),
        ("Der Zeuge Bauer wurde am Montag angehört.", "Bauer"),
        # German records state a maiden name with a bare abbreviation.
        ("Die Kundin, geb. Weber, führt das Konto seit 2011.", "Weber"),
        # An initial before a surname.
        ("B. Winkler hat den Vorgang gezeichnet.", "Winkler"),
    ],
)
def test_anchored_name_patterns_cover_the_shapes_the_hard_strata_exposed(
    analyzer, base_config, text, expected
):
    values = {f.value for f in _findings(analyzer, base_config, text)}
    assert expected in values, f"{expected!r} not caught in {text!r}: {values}"


def test_a_role_noun_anchor_never_swallows_the_honorific(analyzer, base_config):
    """"Der Antragsteller Herr Müller" must yield "Müller". A title inside the
    value keys the pseudonym on the title, so one person becomes two placeholders
    -- the same defect the English-honorific fix closed."""
    values = {f.value for f in _findings(analyzer, base_config, "Der Antragsteller Herr Müller hat gezeichnet.")}
    assert "Müller" in values, values
    assert "Herr Müller" not in values


def test_an_initial_before_a_surname_is_name_shaped():
    """A "Kürzel"/initials column failed the whole-cell shape gate outright,
    because the initial's period read as sentence punctuation."""
    from anonymizer.formats.xlsx_handler import _looks_like_name

    assert _looks_like_name("B. Winkler")
    assert _looks_like_name("M. Schmidt-Rottluff")
    assert not _looks_like_name("b. winkler")  # lowercase is not a name
    assert not _looks_like_name("Das ist ein ganzer Satz.")


def _infer_cols(tmp_path, analyzer, base_config, columns, validations=()):
    """Builds a one-sheet workbook from {header: [values]} and returns the set of
    column letters the CONTENT-based inference claims as people columns."""
    import openpyxl
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    from anonymizer.formats.xlsx_handler import _inferred_name_columns, _sheet_languages

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Daten"
    for c, (header, values) in enumerate(columns.items(), start=1):
        ws.cell(row=1, column=c, value=header)
        for r, v in enumerate(values, start=2):
            ws.cell(row=r, column=c, value=v)
    for formula, applies in validations:
        dv = DataValidation(type="list", formula1=formula)
        ws.add_data_validation(dv)
        dv.add(applies)
    path = tmp_path / "infer.xlsx"
    wb.save(path)
    wb2 = openpyxl.load_workbook(path)
    cfg = {**base_config, "languages": ["de"]}
    found = _inferred_name_columns(wb2, analyzer, cfg, _sheet_languages(wb2, cfg))
    return {col for sheet, col in found if sheet == "Daten"}, get_column_letter


def test_a_column_of_bare_surnames_is_inferred_from_its_content(tmp_path, analyzer, base_config):
    """The fourth corroboration source. Under a header that says nothing, NER
    catches the foreign and rare surnames and misses the everyday-word German
    ones; the caught minority is evidence about the COLUMN, which rescues the
    rest. Measured on the harness this took the shape from 10% to 100%."""
    cols, _ = _infer_cols(
        tmp_path, analyzer, base_config,
        {"Status": ["Öztürk", "Habermehl", "Müller", "Kowalczyk", "Bauer", "Osterkamp", "Koch"]},
    )
    assert "A" in cols, f"a column of surnames was not inferred: {cols}"


def test_a_controlled_vocabulary_column_is_not_inferred_as_people(tmp_path, analyzer, base_config):
    """Status/phase values are capitalized, name-shaped and (in a lookup sheet)
    all distinct. The REPETITION guard is what separates them from a roster."""
    cols, _ = _infer_cols(
        tmp_path, analyzer, base_config,
        {"Phase": ["Offen", "Offen", "Geklaert", "Abgeschlossen", "Offen", "Geklaert", "Offen"]},
    )
    assert "A" not in cols, f"an enum column was read as people: {cols}"


def test_a_dropdown_source_column_is_never_inferred_as_people(tmp_path, analyzer, base_config):
    """A validation SOURCE list is vocabulary by the author's own declaration,
    however name-shaped and however distinct. Measured: without this the fixture's
    `DB_Setup` sheet was read as a column of people and cost 2 false positives."""
    cols, _ = _infer_cols(
        tmp_path, analyzer, base_config,
        {"Vokabular": ["Idee", "Validierung", "Konzeption", "Rollout", "Pilotierung"]},
        validations=[("$A$2:$A$6", "C2:C200")],
    )
    assert "A" not in cols, f"a dropdown source list was read as people: {cols}"


def test_a_short_column_is_not_enough_to_infer_anything(tmp_path, analyzer, base_config):
    """Three values are not a distribution -- inferring from them would make the
    whole-column claim as fragile as the per-cell guess it replaces."""
    cols, _ = _infer_cols(tmp_path, analyzer, base_config, {"Feld_7": ["Öztürk", "Müller"]})
    assert "A" not in cols, f"inferred from too few values: {cols}"


def test_an_anchor_wins_the_span_against_spacys_flat_confidence(analyzer, base_config):
    """The corroboration model rests entirely on the anchored patterns, and they were
    INVISIBLE wherever spaCy also fired.

    Presidio's EntityRecognizer.remove_duplicates() runs inside analyze(): it drops a
    result contained in a higher-scored result of the same entity type. spaCy reports
    PERSON at a flat 0.85, so an anchor scored below that on the identical span was
    deleted before this codebase ever saw it -- and the group then read as an
    uncorroborated bare guess. Measured consequence: an ANCHORED letter scored 0/5
    under corroboration-only while an unanchored memo scored 4/5, because the memo's
    anchors were uncontested."""
    from anonymizer.core import _GATED_NER_SOURCES

    cfg = {**base_config, "languages": ["de"]}
    for text, value in [
        ("Sehr geehrter Herr Winkler,", "Winkler"),
        ("Kunde: Winkler", "Winkler"),
        ("Der Einreicher Winkler hat gezeichnet.", "Winkler"),
    ]:
        found = {f.value: f.source for f in _findings(analyzer, cfg, text)}
        assert value in found, f"{value!r} not found in {text!r}: {found}"
        assert found[value] not in _GATED_NER_SOURCES, (
            f"anchor lost the span to a bare NER guess in {text!r}: {found}"
        )


def test_the_same_value_inherits_corroboration_across_entity_types():
    """Groups are keyed by (type, value), so one name typed PERSON in one sentence and
    ORGANIZATION in another becomes two groups -- and the second was demoted while the
    identical characters were redacted elsewhere. Measured: "Verteiler: Rechtsabteilung,
    Winkler, Innenrevision" types Winkler ORGANIZATION, costing exactly one occurrence
    on nearly every name in the harness (the uniform "4/5")."""
    from anonymizer import core
    from anonymizer.models import Finding as F

    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            F("PERSON", "Winkler", 0.86, "c", "u1", 0, 7, source="PatternRecognizer"),
            F("ORGANIZATION", "Winkler", 0.85, "c", "u2", 0, 7, source="SpacyRecognizer"),
            # An unrelated ORGANIZATION with no corroborated twin still demotes.
            F("ORGANIZATION", "Datenfeeds", 0.85, "c", "u3", 0, 10, source="SpacyRecognizer"),
        ],
        [TextUnit("u", "x")], cfg,
    )
    kept = {(g.entity_type, g.value) for g in result.all_actionable()}
    assert ("ORGANIZATION", "Winkler") in kept, kept
    assert {g.value for g in result.demoted} == {"Datenfeeds"}, [g.value for g in result.demoted]


def test_a_full_name_inherits_from_its_corroborated_surname(monkeypatch):
    """Both directions are needed. A bare "Winkler" inheriting from "Ayse Winkler" was
    handled; the signature line "Mit freundlichen Gruessen Ayse Winkler" forms its OWN
    group and has to inherit from the corroborated bare "Winkler" -- the other way
    round. That single occurrence was the difference between 4/5 and 5/5 on nearly
    every name in the full-letter stratum.

    PERSON is not corroboration-only in the shipped config, so the rule this pins is
    only reachable with the flip on -- which is exactly the state it has to be correct
    in before that flip can ship. Patched in rather than waiting for it."""
    from anonymizer import core
    from anonymizer.models import Finding as F

    monkeypatch.setattr(
        core, "_CORROBORATION_ONLY_ENTITIES",
        frozenset(core._CORROBORATION_ONLY_ENTITIES | {"PERSON"}),
    )
    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            F("PERSON", "Winkler", 0.86, "c", "u1", 0, 7, source="PatternRecognizer"),
            F("PERSON", "Ayse Winkler", 0.85, "c", "u2", 0, 12, source="SpacyRecognizer"),
            # An unrelated bare guess with no corroborated relative still demotes.
            F("PERSON", "Portfoliobeitrag", 0.85, "c", "u3", 0, 16, source="SpacyRecognizer"),
        ],
        [TextUnit("u", "x")], cfg,
    )
    kept = {g.value for g in result.all_actionable()}
    assert {"Winkler", "Ayse Winkler"} <= kept, kept
    assert {g.value for g in result.demoted} == {"Portfoliobeitrag"}, [g.value for g in result.demoted]


def test_a_cell_holding_several_values_is_split_before_the_name_check():
    """The whole-cell override asks whether the ENTIRE cell is a name, so a cell holding
    a name next to anything else was invisible to it -- measured, "von Bergen; intern
    geprueft" scored 0% under corroboration-only. Splitting also fixes a quieter bug:
    "Winkler; Habermehl" passed the shape check as ONE name, so two people shared a
    single pseudonym."""
    from anonymizer.formats.xlsx_handler import _value_segments

    assert [s for _a, _b, s in _value_segments("von Bergen; intern geprueft")] == [
        "von Bergen", "intern geprueft",
    ]
    assert [s for _a, _b, s in _value_segments("Winkler; Habermehl")] == ["Winkler", "Habermehl"]
    assert [s for _a, _b, s in _value_segments("Winkler | intern")] == ["Winkler", "intern"]
    assert [s for _a, _b, s in _value_segments("A_x001E_B")] == ["A", "B"]
    # No separator -> ONE whole-value segment, so the single-value path is unchanged.
    assert _value_segments("Ayse Winkler") == [(0, 12, "Ayse Winkler")]
    # Offsets must index back into the original value, or applying a finding corrupts
    # the cell.
    v = "von Bergen; intern geprueft"
    for start, end, seg in _value_segments(v):
        assert v[start:end] == seg


def test_a_multi_value_people_cell_claims_only_its_name_segment(tmp_path, analyzer, base_config):
    """End-to-end through the real pipeline: the name is claimed, the commentary next to
    it is not, and the span is clean (no trailing separator)."""
    import openpyxl

    from anonymizer.pipeline import scan_document

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Daten"
    ws.cell(row=1, column=1, value="Vermerk")
    for r, v in enumerate(
        ["Öztürk; intern geprüft", "Habermehl; intern geprüft", "Müller; intern geprüft",
         "Kowalczyk; intern geprüft", "Bauer; intern geprüft", "Osterkamp; intern geprüft"],
        start=2,
    ):
        ws.cell(row=r, column=1, value=v)
    path = tmp_path / "multi.xlsx"
    wb.save(path)

    values = {g.value for g in scan_document(path, analyzer, base_config).all_actionable()}
    assert "Müller" in values, f"the everyday-word surname leaked: {values}"
    assert not any("geprüft" in v for v in values), f"commentary claimed as a name: {values}"
    assert not any(v.endswith(";") or v.endswith("|") for v in values), f"dirty span: {values}"


@pytest.mark.parametrize("text", ["KUNDE: WINKLER", "kunde: Winkler", "Kunde: Winkler"])
def test_a_shouted_label_still_anchors_the_name(analyzer, base_config, text):
    """Forms and legacy exports shout their labels. With the whole anchored pattern
    case-sensitive, "KUNDE: WINKLER" matched nothing -- measured at 30% for German
    common-noun surnames against 100% for the ordinary "Kunde:" form. The flag is
    SCOPED to the label alternation: a blanket IGNORECASE would let the name part match
    ordinary lowercase German words, which is the bug BIC_CODE already documents."""
    values = {f.value.lower() for f in _findings(analyzer, base_config, text)}
    assert any("winkler" in v for v in values), f"label did not anchor in {text!r}: {values}"


def test_a_known_name_is_found_inside_an_underscore_identifier():
    """An underscore is a word character, so the `\w` boundary propagation used could
    never see a name inside the identifiers a bank's systems generate. Measured: the
    hyphen and UNC-path forms already propagated (those separators are non-word) while
    "AKTE_Winkler_2024" and "Vertrag_Winkler_final_v2.pdf" silently did not."""
    from anonymizer.core import _compiled_propagate_patterns

    (_et, _v, rx), = _compiled_propagate_patterns((("PERSON", "Winkler"),))
    for haystack in [
        "AKTE_Winkler_2024",
        "Vertrag_Winkler_final_v2.pdf",
        "K-Winkler-2024",
        r"\fileserver\Kunden\Winkler\2024",
        "Winkler",
    ]:
        assert rx.search(haystack), f"known name not reached inside {haystack!r}"
    # The alphanumeric half of the guard is unchanged -- a common-noun surname must
    # still refuse to match inside a longer word.
    for haystack in ["Bergstraße", "Winklerhof"]:
        (_et2, _v2, rx2), = _compiled_propagate_patterns((("PERSON", "Berg" if "Berg" in haystack else "Winkler"),))
        assert not rx2.search(haystack), f"over-matched inside {haystack!r}"


# --- 2026-07-30: Art. 9 disclosure FRAMES ------------------------------------
# The mechanism the word lists structurally cannot be. A frame keys on the
# SENTENCE SHAPE, so it generalizes to conditions and bodies nobody listed.


@pytest.mark.parametrize("text,needle", [
    # HELD OUT of every shipped word list on purpose -- if these are caught, the
    # frame generalizes rather than the list containing its own benchmark.
    ("Der Kunde leidet an einer chronischen Sarkoidose.", "Sarkoidose"),
    ("Die Mitarbeiterin ist im Mai an Tuberkulose erkrankt.", "Tuberkulose"),
    ("Sie ist nachts auf das Beatmungsgeraet angewiesen.", "Beatmungsgeraet"),
    ("Er ist krankgeschrieben wegen einer Kniearthroskopie.", "Kniearthroskopie"),
    ("Er konvertierte im Jahr 2018 zum Buddhismus.", "Buddhismus"),
    ("Sie wurde in den Wirtschaftsausschuss gewaehlt.", "Wirtschaftsausschuss"),
])
def test_an_art9_frame_catches_a_value_no_word_list_contains(analyzer, base_config, text, needle):
    found = {f.value for f in _art9(analyzer, base_config, text)}
    assert any(needle in v for v in found), f"{needle!r} leaked from {text!r}: {found}"


@pytest.mark.parametrize("text,needle", [
    # The half a frame CANNOT reach: "besucht die Moschee" and "besucht die
    # Filiale" are the same shape, so only vocabulary separates them.
    ("Er besucht jeden Freitag die Moschee in der Innenstadt.", "Moschee"),
    ("Waehrend des Ramadan bittet sie um spaetere Termine.", "Ramadan"),
    ("Er wurde in den Betriebsrat gewaehlt und ist freigestellt.", "Betriebsrat"),
    ("Fuer die Ausfalltage wurde Streikgeld ausgezahlt.", "Streikgeld"),
    ("Die Familie kam 1994 als Kontingentfluechtlinge nach Deutschland.", "Kontingentfl"),
    ("Das Beratungsgespraech wurde auf Romanes gefuehrt.", "Romanes"),
    ("Der Kunde ist dauerhaft auf den Rollstuhl angewiesen.", "Rollstuhl"),
])
def test_art9_vocabulary_covers_what_a_frame_cannot(analyzer, base_config, text, needle):
    found = {f.value for f in _art9(analyzer, base_config, text)}
    assert any(needle in v for v in found), f"{needle!r} leaked from {text!r}: {found}"


@pytest.mark.parametrize("text", [
    # Every "...gemeinde" term was tried in the religion list and REMOVED: they sit
    # uninflected inside organisation names, which breaks the property that makes
    # these lists safe to extend -- they match uninflected forms only, so \b fails
    # on a trailing inflection and org names survive. This pins that they stay out.
    "Neuapostolische Kirchengemeinde Frankfurt",
    "Evangelische Freikirche Bonn e.V.",
    "Katholische Pfarrgemeinde St. Martin",
])
def test_religion_vocabulary_still_leaves_organisation_names_intact(analyzer, base_config, text):
    found = {f.value for f in _art9(analyzer, base_config, text)}
    assert not found, f"an organisation name was claimed as Art. 9 data: {found}"


def test_art9_frames_stay_in_the_review_tier(analyzer, base_config):
    """These entity types carry a one-way `anonymize` action. A frame match is
    strong evidence but not proof, and one-way destruction is unrecoverable -- so a
    frame must never reach the auto-accept tier. The reviewer decides."""
    findings = _art9(analyzer, base_config, "Der Kunde leidet an einer chronischen Sarkoidose.")
    assert findings
    assert all(f.score < base_config["tiers"]["high"] for f in findings), [
        (f.value, f.score) for f in findings
    ]


def test_same_sex_partnership_is_detected_from_the_possessive_alone(analyzer, base_config):
    """German marks it in the POSSESSIVE: "seine Ehefrau" is an ordinary man's wife,
    "ihre Ehefrau" is a woman's wife. That pronoun is the entire Art. 9 signal, which
    is exactly why a bare "Ehefrau" must never be listed -- it would flag every married
    customer in the file and destroy the word one-way."""
    found = {f.value for f in _art9(analyzer, base_config, "Sie lebt mit ihrer Ehefrau in Koeln.")}
    assert any("Ehefrau" in v for v in found), f"same-sex partnership leaked: {found}"

    ordinary = {f.value for f in _art9(analyzer, base_config, "Er hat das Konto mit seiner Ehefrau eroeffnet.")}
    assert not any("Ehefrau" in v for v in ordinary), f"an ordinary marriage was claimed: {ordinary}"


def test_a_party_is_detected_mid_sentence_with_a_lowercase_article(analyzer, base_config):
    """A party is named mid-sentence far more often than at its start. With a
    capital-only "Die Gruenen" the most ordinary phrasing there is went unmatched, and
    political_party measured 0%."""
    found = {f.value for f in _art9(analyzer, base_config, "Er kandidierte bei der Kommunalwahl fuer die Gruenen.")}
    assert any("rünen" in v or "ruenen" in v for v in found), f"party affiliation leaked: {found}"


# --- 2026-07-30: OCR-damaged names -------------------------------------------


@pytest.mark.parametrize("damaged,clean", [
    ("Wlnkler", "Winkler"),        # i -> l
    ("Mul1er", "Müller"),          # umlaut dropped, l -> 1
    ("Miiller", "Müller"),         # the classic OCR umlaut
    ("0sterkamp", "Osterkamp"),    # O -> 0
    ("Kretschrnar", "Kretschmar"), # m -> rn
    ("Mueller", "Müller"),         # transliterated umlaut
])
def test_an_ocr_skeleton_folds_the_damage_a_scanner_actually_does(damaged, clean):
    from anonymizer.core import ocr_skeleton

    assert ocr_skeleton(damaged) == ocr_skeleton(clean), (
        f"{damaged!r} and {clean!r} fold differently: "
        f"{ocr_skeleton(damaged)!r} vs {ocr_skeleton(clean)!r}"
    )


@pytest.mark.parametrize("a,b", [
    # b/h is deliberately NOT folded: it would make these the same skeleton and
    # both are real German surnames -- the one collision in the table that could
    # redact the wrong person's name.
    ("Bauer", "Hauer"),
    ("Koch", "Kock"),
    ("Berg", "Burg"),
    ("Fischer", "Fisher"),
])
def test_the_ocr_skeleton_does_not_collide_real_surnames(a, b):
    from anonymizer.core import ocr_skeleton

    assert ocr_skeleton(a) != ocr_skeleton(b), f"{a!r} and {b!r} collided on {ocr_skeleton(a)!r}"


def test_an_ocr_damaged_name_inherits_its_clean_twins_corroboration():
    """A scanner mangles SOME occurrences and not others, so "Mul1er" forms its own
    group beside a corroborated "Mueller" and was demoted -- leaving the surname
    legible in exactly the documents where it was hardest to read.

    Note this is INHERITANCE, not a new source: tagging the OCR match "propagation"
    (the first attempt) rescued nothing, because propagation is deliberately not
    corroboration -- so the damaged group was still demoted."""
    from anonymizer import core
    from anonymizer.models import Finding as F

    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            F("PERSON", "Mueller", 0.86, "c", "u1", 0, 7, source="PatternRecognizer"),
            F("NER_MISC", "Mul1er", 0.85, "c", "u2", 0, 6, source="SpacyRecognizer"),
            # An unrelated damaged-looking token with no corroborated twin still demotes.
            F("NER_MISC", "Datenfeeds", 0.85, "c", "u3", 0, 10, source="SpacyRecognizer"),
        ],
        [TextUnit("u", "x")], cfg,
    )
    kept = {g.value for g in result.all_actionable()}
    assert "Mul1er" in kept, f"the OCR-damaged spelling leaked: {kept}"
    assert {g.value for g in result.demoted} == {"Datenfeeds"}, [g.value for g in result.demoted]


def test_an_ambiguous_ocr_skeleton_is_never_inherited():
    """If two DIFFERENT corroborated names fold to the same skeleton, nothing can
    tell which one a damaged token was -- so it must not be guessed. Redacting a
    person under another person's pseudonym is worse than surfacing for review."""
    from anonymizer import core
    from anonymizer.models import Finding as F

    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            F("PERSON", "Muller", 0.86, "c", "u1", 0, 6, source="PatternRecognizer"),
            F("PERSON", "Müller", 0.86, "c", "u2", 0, 6, source="PatternRecognizer"),
            F("NER_MISC", "Mul1er", 0.85, "c", "u3", 0, 6, source="SpacyRecognizer"),
        ],
        [TextUnit("u", "x")], cfg,
    )
    assert "Mul1er" in {g.value for g in result.demoted}, [g.value for g in result.demoted]


# --- 2026-07-30: proprietary product / project names --------------------------


def test_a_listed_product_name_corroborates_whatever_the_model_guessed(analyzer, base_config):
    """Measured on the audit fixture: every tool name that appeared once in a column
    headed "Eingesetztes Tool" was learned from the document's own structure and then
    found everywhere. The only two that LEAKED (Alteryx, OpenClaw) were the two that
    appear ONLY in prose, with no declaring column to learn them from -- and the model
    typed both LOCATION, a bare guess that corroboration-only demotes.

    Note this is corroboration across ALL the NER types, not one: the model has no idea
    what kind of thing a proprietary name is and guessed LOCATION here, ORGANIZATION and
    MISC for other tools in the same file.

    CORROBORATION, NOT DETECTION -- and the distinction is load-bearing here. The
    gazetteer never creates a finding, so in a short sentence where the model produces
    no candidate at all there is nothing for it to confirm; what it fixes is the case
    where a candidate EXISTS but reads as a bare guess.
    """
    from anonymizer.core import _GATED_NER_SOURCES, PRODUCT_NAME_SOURCE, is_known_product
    from anonymizer.models import Finding as F

    assert is_known_product("Alteryx")
    assert is_known_product("alteryx"), "matching must be case-insensitive"
    assert is_known_product("Kondor+"), "punctuation in a product name must survive"
    assert not is_known_product("Portfoliobeitrag")
    # The source it re-stamps must actually count as corroboration.
    assert PRODUCT_NAME_SOURCE not in _GATED_NER_SOURCES

    # A LOCATION guess on a listed product survives corroboration-only; an
    # unlisted one is still demoted.
    from anonymizer import core

    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            F("LOCATION", "Alteryx", 0.85, "c", "u1", 0, 7, source=PRODUCT_NAME_SOURCE),
            F("LOCATION", "Datenfeeds", 0.85, "c", "u2", 0, 10, source="SpacyRecognizer"),
        ],
        [TextUnit("u", "x")], cfg,
    )
    assert "Alteryx" in {g.value for g in result.all_actionable()}
    assert {g.value for g in result.demoted} == {"Datenfeeds"}


def test_the_project_list_is_what_catches_an_internal_codename(monkeypatch):
    """Codenames are chosen to be unremarkable, so they are ordinary German words --
    "Nordstern", "Seidenpfad", "Habicht", "Delphin" are all in the audit fixture and all
    normal nouns. No shipped list and no heuristic can find those; project_names.txt is
    the mechanism, which is why it ships empty rather than absent."""
    from anonymizer import core

    # The shipped list cannot contain an internal codename, by construction.
    assert not core.is_known_product("OpenClaw")

    # Adding it to the user's own list is what finds it.
    monkeypatch.setattr(core, "_product_names", lambda: frozenset({"openclaw", "nordstern"}))
    assert core.is_known_product("OpenClaw")
    assert core.is_known_product("nordstern")


def test_a_missing_product_list_degrades_instead_of_failing(monkeypatch, tmp_path):
    """Same contract as the given-name list: a detection INPUT that can hard-fail a scan
    by being absent would be a worse bug than the recall it buys."""
    from anonymizer import core

    core._product_names.cache_clear()
    monkeypatch.setattr(core, "_PRODUCT_NAMES_PATHS", (tmp_path / "nope.txt",))
    try:
        assert core._product_names() == frozenset()
        assert not core.is_known_product("Alteryx")
    finally:
        core._product_names.cache_clear()  # the real list must be reloaded for other tests


@pytest.mark.parametrize("word,expected", [
    # Proprietary: neither German vocabulary nor a German compound.
    ("Alteryx", True), ("Signavio", True), ("Camunda", True), ("Collibra", True),
    # German COMPOUNDS are absent from any word list while being entirely ordinary.
    # Three-part splitting is what gets Marktdatengrundlage right.
    ("Portfoliobeitrag", False), ("Marktdatengrundlage", False), ("Bearbeitungszeit", False),
    ("Hauptzielgruppe", False), ("Projektlaufzeit", False),
    # Nominalizer suffixes are German morphology however absent from a vector table.
    ("Derivatefreiheit", False), ("Effizienz", False), ("Reaktionszeiten", False),
    # ALL-CAPS is an acronym or field code, not a product-name shape.
    ("CAPEX", False), ("OPEX", False),
    # ASCII-transliterated umlauts are ordinary German business text.
    ("Geschaeftsbedingungen", False),
])
def test_proprietary_name_heuristic_separates_products_from_german_compounds(
    analyzer, word, expected
):
    from anonymizer.core import looks_like_proprietary_name

    vocab = analyzer.nlp_engine.nlp["de"].vocab
    known = lambda w: vocab[w].has_vector  # noqa: E731
    assert looks_like_proprietary_name(word, known) is expected, word


def test_a_proprietary_candidate_stays_in_the_demoted_band():
    """The weakest signal the tool produces, so it must never reach the list the tool
    acts on. Measured without the guard: the exact-value inheritance rule promoted these
    and they cost 3 decoy false positives."""
    from anonymizer import core
    from anonymizer.models import Finding as F

    cfg = {"entities": {}, "tiers": {"high": 0.9, "medium": 0.5}, "corroboration_only": True}
    result = core.build_scan_result(
        [
            # The same value corroborated elsewhere would normally promote the group.
            F("PERSON", "Zephyrix", 0.86, "c", "u1", 0, 8, source="PatternRecognizer"),
            F("NER_MISC", "Zephyrix", 0.85, "c", "u2", 0, 8,
              source=core.PROPRIETARY_CANDIDATE_SOURCE),
        ],
        [TextUnit("u", "x")], cfg,
    )
    assert "Zephyrix" in {g.value for g in result.demoted}, [g.value for g in result.demoted]
    assert all(g.entity_type != "NER_MISC" for g in result.all_actionable())
