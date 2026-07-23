"""Phase-4 fail-loud coverage: numeric-cell detection (#4), metadata scrubbing
(#3), and the recognizer-independent literal-residual backstop (#5)."""

import zipfile

import openpyxl
import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

from anonymizer.formats import docx_handler
from anonymizer.pipeline import _literal_residual, _output_text_blob, apply_document, scan_document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
IBAN = "DE89370400440532013000"


def _raw_package_text(path) -> str:
    """Every byte of every part in an OOXML package, decoded. The strongest,
    completely handler-INDEPENDENT proof that a value really left the file --
    a re-scan through the same handler that wrote the file cannot give that."""
    with zipfile.ZipFile(path) as zf:
        return "\n".join(zf.read(n).decode("utf-8", "ignore") for n in zf.namelist())


def _append_xml(doc, xml: str) -> None:
    doc.element.body.append(etree.fromstring(xml))


def _add_textbox(doc, text: str) -> None:
    """python-docx has no text-box API, so inject a minimal VML text box."""
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}" xmlns:v="urn:schemas-microsoft-com:vml"><w:r><w:pict><v:shape>'
        f"<v:textbox><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        f"</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>",
    )


def _add_hyperlink_paragraph(doc, text: str) -> None:
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}"><w:hyperlink><w:r><w:t>{text}</w:t></w:r></w:hyperlink></w:p>',
    )


def test_numeric_account_cell_detected_and_redacted(tmp_path, analyzer, base_config, mapping_db_path):
    """An account number stored as a NUMBER (not text) used to be invisible to
    scan and verify -> emitted in a 'verified' file. It must now be detected via
    the column header context and redacted."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Kontonummer"  # header -> supplies context for the bare number
    ws["A2"] = 1234567890  # stored as an int, data_type == "n"
    path = tmp_path / "accounts.xlsx"
    wb.save(path)

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    assert any("1234567890" in g.value for g in grouped), "numeric account number must be detected"

    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)
    assert "1234567890" not in _output_text_blob(out_path)
    assert _literal_residual(out_path, ["1234567890"]) == []


def test_document_metadata_is_scrubbed(tmp_path, analyzer, base_config, mapping_db_path):
    """The author / last-modified-by carry the real name but are never in body
    text; they must be blanked, and the literal backstop must confirm the name
    is gone from the whole file."""
    doc = Document()
    doc.add_paragraph("Kunde: Hans Mueller")
    doc.core_properties.author = "Hans Mueller"
    doc.core_properties.last_modified_by = "Hans Mueller"
    path = tmp_path / "letter.docx"
    doc.save(path)

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)

    out_doc = Document(out_path)
    assert out_doc.core_properties.author == ""
    assert out_doc.core_properties.last_modified_by == ""
    assert _literal_residual(out_path, ["Hans Mueller"]) == []


def test_textbox_text_is_scanned_and_redacted(tmp_path, analyzer, base_config, mapping_db_path):
    """PII inside a Word text box (w:txbxContent) was invisible to scan AND to
    the output re-scan -- a false-clean leak in letterhead/form templates."""
    doc = Document()
    doc.add_paragraph("Vertrag mit der Musterbank.")
    _add_textbox(doc, f"Zahlungen an IBAN {IBAN}")
    path = tmp_path / "textbox.docx"
    doc.save(path)

    units = docx_handler.extract_text_units(path)
    assert any(IBAN in u.text for u in units), "text-box text must be extracted"

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    assert any(IBAN in g.value.replace(" ", "") for g in grouped)
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)
    assert IBAN not in _output_text_blob(out_path).replace(" ", "")


# --- B1: channels the handler-dependent verify never looks at ----------------
# verify_output re-scans the OUTPUT through the SAME handler that wrote it, and
# only re-checks values that were already DECIDED. PII in a location the
# extractor never reaches is therefore in neither set: it is "verified" without
# ever being looked at, and apply reports success. Each test below plants PII in
# exactly one such location, with NO copy of it in the body, so the value can
# never enter `removed_values` -- the false-clean condition.


def test_core_properties_title_pii_does_not_survive(tmp_path, analyzer, base_config, mapping_db_path):
    """FALSE-CLEAN (B1a/B3): docProps/core.xml title/subject/description/keywords
    are body-text-invisible, so a customer name in "Kreditakte Hans Mueller" --
    the classic place for one -- was never scanned, never removed, and never
    re-checked, yet apply reported the file verified."""
    doc = Document()
    doc.add_paragraph("Anlage zum Vertrag.")
    doc.core_properties.title = "Kreditakte Hans Mueller"
    doc.core_properties.subject = f"Konto {IBAN}"
    doc.core_properties.keywords = "Hans Mueller"
    doc.core_properties.comments = "Beschwerde von Hans Mueller"  # dc:description
    doc.core_properties.category = "Hans Mueller"
    path = tmp_path / "titled.docx"
    doc.save(path)

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)

    blob = _raw_package_text(out_path)
    assert "Hans Mueller" not in blob, "core.xml title/keywords/description leaked the customer name"
    assert IBAN not in blob, "core.xml subject leaked the IBAN"


def _inject_custom_props(path, name: str, value: str) -> None:
    """Adds a real docProps/custom.xml part (python-docx has no API for it):
    the part, its content-type override and its package relationship."""
    ct = "application/vnd.openxmlformats-officedocument.custom-properties+xml"
    custom = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        f'<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="2" name="{name}">'
        f"<vt:lpwstr>{value}</vt:lpwstr></property></Properties>"
    )
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        contents = {n: zf.read(n) for n in names}
    contents["[Content_Types].xml"] = contents["[Content_Types].xml"].replace(
        b"</Types>", f'<Override PartName="/docProps/custom.xml" ContentType="{ct}"/></Types>'.encode()
    )
    contents["_rels/.rels"] = contents["_rels/.rels"].replace(
        b"</Relationships>",
        b'<Relationship Id="rIdCustom" Type="http://schemas.openxmlformats.org/officeDocument/2006/'
        b'relationships/custom-properties" Target="docProps/custom.xml"/></Relationships>',
    )
    contents["docProps/custom.xml"] = custom.encode("utf-8")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for n in names + ["docProps/custom.xml"]:
            zf.writestr(n, contents[n])


def test_custom_document_properties_pii_does_not_survive(tmp_path, analyzer, base_config, mapping_db_path):
    """FALSE-CLEAN (B1a/B3): docProps/custom.xml holds arbitrary user-defined
    fields -- bank templates put the case owner / customer there. Nothing in the
    pipeline read or scrubbed that part."""
    doc = Document()
    doc.add_paragraph("Anlage zum Vertrag.")
    path = tmp_path / "custom.docx"
    doc.save(path)
    _inject_custom_props(path, "Kundenbetreuer", "Hans Mueller")

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)

    assert "Hans Mueller" not in _raw_package_text(out_path), "docProps/custom.xml leaked the name"


def test_hyperlink_target_url_does_not_survive(tmp_path, analyzer, base_config, mapping_db_path):
    """FALSE-CLEAN (B1b): only the hyperlink DISPLAY text was ever scanned. The
    TARGET lives in a relationship ATTRIBUTE (word/_rels/document.xml.rels), which
    neither the extractor nor the literal backstop (element text only) ever read --
    so mailto:hans.mueller@bank.de shipped inside a "verified" document."""
    doc = Document()
    rid = doc.part.relate_to("mailto:hans.mueller@bank.de", RT.HYPERLINK, is_external=True)
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}" xmlns:r="{R_NS}"><w:hyperlink r:id="{rid}">'
        f"<w:r><w:t>Kontakt aufnehmen</w:t></w:r></w:hyperlink></w:p>",
    )
    path = tmp_path / "target.docx"
    doc.save(path)

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    assert any("hans.mueller@bank.de" in g.value for g in grouped), "hyperlink target must be surfaced for review"
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)

    assert "hans.mueller@bank.de" not in _raw_package_text(out_path), "hyperlink target leaked"


def test_literal_backstop_reads_hyperlink_targets(tmp_path):
    """The backstop is the LAST line of defence, so it must be handler-INDEPENDENT.
    A relationship Target is an XML ATTRIBUTE: an itertext() sweep of the package
    could not see it, so a removed value surviving there passed verification."""
    doc = Document()
    doc.add_paragraph("Text.")
    doc.part.relate_to("mailto:hans.mueller@bank.de", RT.HYPERLINK, is_external=True)
    path = tmp_path / "rel.docx"
    doc.save(path)

    assert _literal_residual(path, ["hans.mueller@bank.de"]) == ["hans.mueller@bank.de"]
    # Internal targets are package paths, never PII -- including them would only
    # manufacture phantom substring matches and spurious hard-fails.
    assert _literal_residual(path, ["theme1.xml"]) == []


def test_literal_backstop_descends_into_embedded_packages(tmp_path):
    """A chart's source workbook is a whole OOXML package nested inside an already-
    compressed part, so its text is invisible both to an XML sweep of the outer
    parts and to a raw byte scan of the file. A removed value surviving only there
    must still fail loud."""
    import io

    doc = Document()
    doc.add_paragraph("Text.")
    path = tmp_path / "embed.docx"
    doc.save(path)

    inner = openpyxl.Workbook()
    inner.active["A1"] = "Hans Mueller"
    buf = io.BytesIO()
    inner.save(buf)
    with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/embeddings/Microsoft_Excel_Sheet1.xlsx", buf.getvalue())

    assert path.read_bytes().find(b"Hans Mueller") == -1, "precondition: not findable by a raw byte scan"
    assert _literal_residual(path, ["Hans Mueller"]) == ["Hans Mueller"]


def test_literal_backstop_reads_every_xml_attribute(tmp_path):
    """FALSE-CLEAN (B1): the backstop read element TEXT plus exactly ONE attribute
    (an External relationship Target). EVERY other XML attribute was invisible to
    BOTH halves of the gate, so a value the reviewer explicitly DECIDED AND REMOVED
    could survive verbatim in one and the file was still committed as verified.
    Three real carriers, one per format family."""
    from openpyxl.worksheet.hyperlink import Hyperlink

    # (1) an Excel cell hyperlink TOOLTIP -- xl/worksheets/sheet1.xml @tooltip
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Kontakt"
    ws["A1"].hyperlink = Hyperlink(ref="A1", target="https://intranet/", tooltip="Hans Mueller anrufen")
    tip_path = tmp_path / "tooltip.xlsx"
    wb.save(tip_path)
    assert _literal_residual(tip_path, ["Hans Mueller"]) == ["Hans Mueller"]

    # (2) an Excel SHEET NAME -- xl/workbook.xml <sheet name="...">
    wb2 = openpyxl.Workbook()
    wb2.active.title = "Hans Mueller"
    wb2.active["A1"] = "Umsatz"
    sheet_path = tmp_path / "sheetname.xlsx"
    wb2.save(sheet_path)
    assert _literal_residual(sheet_path, ["Hans Mueller"]) == ["Hans Mueller"]

    # (3) a Word FIELD-CODE hyperlink -- w:fldSimple/@w:instr in word/document.xml
    doc = Document()
    doc.add_paragraph("Text.")
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}"><w:fldSimple w:instr=\' HYPERLINK "mailto:hans.mueller@bank.de" \'>'
        f"<w:r><w:t>Kontakt</w:t></w:r></w:fldSimple></w:p>",
    )
    fld_path = tmp_path / "fld.docx"
    doc.save(fld_path)
    assert _literal_residual(fld_path, ["hans.mueller@bank.de"]) == ["hans.mueller@bank.de"]

    # ...and the package plumbing must still not manufacture phantom hard-fails.
    assert _literal_residual(fld_path, ["theme1.xml"]) == []


def test_docx_field_code_hyperlink_is_surfaced_and_redacted(tmp_path, analyzer, base_config, mapping_db_path):
    """FALSE-CLEAN (B1b): a Word FIELD CODE (w:instrText / w:fldSimple/@w:instr) is
    neither a w:t run nor a relationship, so python-docx never yielded it and the
    scan never surfaced the URL -- it entered no decision set, so the literal
    backstop had nothing to check either. Mail-merge and template documents (the
    common bank case) produce exactly this. The document's ONLY PII is the field
    code, so nothing else can pull the address into `removed_values`."""
    doc = Document()
    doc.add_paragraph("Bitte den Ansprechpartner kontaktieren.")
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}"><w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r><w:instrText xml:space="preserve"> HYPERLINK "mailto:hans.mueller@bank.de" </w:instrText></w:r>'
        f'<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        f"<w:r><w:t>Kontakt</w:t></w:r>"
        f'<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>',
    )
    path = tmp_path / "fieldcode.docx"
    doc.save(path)

    units = docx_handler.extract_text_units(path)
    assert any("hans.mueller@bank.de" in u.text for u in units), "field-code text must be extracted"

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    assert any("hans.mueller@bank.de" in g.value for g in grouped), "field-code URL must be surfaced for review"
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)
    assert "hans.mueller@bank.de" not in _raw_package_text(out_path), "field-code hyperlink leaked"


def test_docx_fld_simple_instruction_is_surfaced_and_redacted(tmp_path, analyzer, base_config, mapping_db_path):
    """FALSE-CLEAN (B1b): the same hole in the w:fldSimple form, where the whole
    instruction -- including a DOCPROPERTY payload naming the customer -- lives in
    an ATTRIBUTE that no text sweep reads."""
    doc = Document()
    doc.add_paragraph("Anlage.")
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}"><w:fldSimple w:instr=\' DOCPROPERTY "Kunde Hans Mueller" \'>'
        f"<w:r><w:t>Kunde</w:t></w:r></w:fldSimple></w:p>",
    )
    path = tmp_path / "fldsimple.docx"
    doc.save(path)

    units = docx_handler.extract_text_units(path)
    assert any("Hans Mueller" in u.text for u in units), "w:fldSimple instruction must be extracted"

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)
    assert "Hans Mueller" not in _raw_package_text(out_path), "w:fldSimple instruction leaked"


def test_unreadable_package_fails_loud_instead_of_skipping_silently(tmp_path):
    """FAIL-LOUD (B1c): the auxiliary pass swallowed BadZipFile/OSError and any
    exception from an embedded workbook, returning nothing. Input documents are
    UNTRUSTED, so "cannot read it" is exactly the case that must NOT pass: PII
    living only inside an unreadable embedding was neither surfaced nor checked and
    shipped silently, while the comment claimed the backstop still covered it."""
    from anonymizer.formats import run_replace
    from anonymizer.models import ProcessingError

    broken = tmp_path / "broken.docx"
    broken.write_bytes(b"this is not a zip at all")
    with pytest.raises(ProcessingError):
        run_replace.aux_text_units(broken)
    with pytest.raises(ProcessingError):
        run_replace.apply_aux_parts(broken, None, {}, {}, None)

    # A perfectly valid package carrying a CORRUPT embedded workbook: the outer
    # file opens, so nothing else in the pipeline notices.
    doc = Document()
    doc.add_paragraph("Text.")
    path = tmp_path / "embed.docx"
    doc.save(path)
    with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/embeddings/Microsoft_Excel_Sheet1.xlsx", b"corrupt bytes")
    with pytest.raises(ProcessingError):
        run_replace.aux_text_units(path)
    with pytest.raises(ProcessingError):
        run_replace.apply_aux_parts(path, None, {}, {}, None)


def test_rewritten_hyperlink_target_stays_a_valid_uri(tmp_path, analyzer, base_config, mapping_db_path):
    """CORRUPTION RISK (B1b-uri): a hyperlink Target that is entirely PII was
    rewritten to a bare "[EMAIL]" -- the URI scheme dropped, and square brackets
    inserted, which RFC 3986 reserves for IPv6 literals. Word treats an external
    relationship with an invalid Target as unreadable content and offers to
    "repair" the file, so for this tool that means the user's colleague receives a
    corrupt document. The rewritten target must stay scheme-preserving and valid."""
    doc = Document()
    rid = doc.part.relate_to("mailto:hans.mueller@bank.de", RT.HYPERLINK, is_external=True)
    _append_xml(
        doc,
        f'<w:p xmlns:w="{W_NS}" xmlns:r="{R_NS}"><w:hyperlink r:id="{rid}">'
        f"<w:r><w:t>Kontakt</w:t></w:r></w:hyperlink></w:p>",
    )
    path = tmp_path / "uri.docx"
    doc.save(path)

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)

    with zipfile.ZipFile(out_path) as zf:
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    tree = etree.fromstring(rels.encode("utf-8"))
    targets = [el.get("Target") for el in tree.iter() if el.get("TargetMode") == "External"]
    assert targets, "the external relationship must survive -- only its PII is removed"
    for target in targets:
        assert "[" not in target and "]" not in target, f"RFC 3986 reserves [ ] : {target}"
        assert target.startswith("mailto:"), f"URI scheme dropped: {target}"


def test_safe_relationship_target_preserves_scheme_and_escapes_brackets():
    """Unit-level guard for the rewrite above, including the case the integration
    test cannot stage: a redaction that consumed the WHOLE target, scheme and all."""
    from anonymizer.formats.run_replace import _safe_relationship_target

    assert _safe_relationship_target("mailto:a@b.de", "mailto:[EMAIL]") == "mailto:%5BEMAIL%5D"
    assert _safe_relationship_target("https://intra/x", "[PERSON_1]") == "https://%5BPERSON_1%5D"
    assert _safe_relationship_target("https://intra/x", "https://intra/x") == "https://intra/x"


def test_hyperlink_text_is_scanned_and_redacted(tmp_path, analyzer, base_config, mapping_db_path):
    """python-docx's p.runs skips runs nested in w:hyperlink, so PII in link
    display text was never scanned or replaced."""
    doc = Document()
    _add_hyperlink_paragraph(doc, f"Konto {IBAN} ansehen")
    path = tmp_path / "link.docx"
    doc.save(path)

    units = docx_handler.extract_text_units(path)
    assert any(IBAN in u.text for u in units), "hyperlink text must be extracted"

    grouped = scan_document(path, analyzer, base_config).all_actionable()
    for g in grouped:
        g.action = "anonymize"
    out_path, _ = apply_document(path, grouped, analyzer, base_config, mapping_db_path)
    assert IBAN not in _output_text_blob(out_path).replace(" ", "")
