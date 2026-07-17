from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from ..core import detect_unit
from ..models import TextUnit
from .run_replace import XmlRunAdapter, apply_findings_to_runs, runs_text_and_spans

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
EXTRA_PARTS = ["word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"]

EXTENSIONS = (".docx", ".doc")


def _para_run_elements(p_elem):
    """The runs of a paragraph in document order: direct `w:r` children PLUS
    runs nested in a `w:hyperlink` (python-docx's `p.runs` skips those, so PII
    in hyperlink display text was never scanned or redacted).

    Deliberately does NOT descend into the paragraph's descendants: a run that
    holds a drawing keeps its own (usually empty) text here, and the paragraphs
    inside that drawing's text box are yielded separately by
    `_textbox_paragraphs`. Descending would redact text-box content twice."""
    out = []
    for child in p_elem:
        if child.tag == f"{W}r":
            out.append(child)
        elif child.tag == f"{W}hyperlink":
            out.extend(child.findall(f"{W}r"))
    return out


def paragraph_runs(p) -> list:
    """Run objects for a paragraph, matching `_para_run_elements`. Detection and
    replacement BOTH build their text from this one list, so their coordinate
    systems are identical by construction (no offset drift)."""
    return [Run(r, p) for r in _para_run_elements(p._p)]


def paragraph_text(p) -> str:
    return runs_text_and_spans(paragraph_runs(p))[0]


def _textbox_paragraphs(doc: Document):
    """Paragraphs inside drawing/VML text boxes (`w:txbxContent`), in the body
    and in every header/footer. `doc.paragraphs` never returns these, so PII in
    a letterhead or form text box -- common in bank templates -- was invisible
    to both scan and the output re-scan."""
    roots = [doc.element.body]
    for section in doc.sections:
        for container in (section.header, section.footer):
            roots.append(container._element)
    for root in roots:
        for txbx in root.iter(f"{W}txbxContent"):
            for p_elem in txbx.iter(f"{W}p"):
                yield Paragraph(p_elem, doc)


def _iter_table_paragraphs(table):
    """A table's cell paragraphs, RECURSING into nested tables (a table inside a
    cell). `doc.tables` returns only top-level tables and `cell.paragraphs` does
    not descend into a nested table, so tables-within-tables -- common in bank
    form layouts -- were scanned and redacted by neither path (a silent leak)."""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                yield from _iter_table_paragraphs(table)
    yield from _textbox_paragraphs(doc)


def _extra_parts_paragraphs(path: Path):
    with zipfile.ZipFile(path, "r") as zf:
        names = zf.namelist()
        contents = {part: zf.read(part) for part in EXTRA_PARTS if part in names}
    for part, blob in contents.items():
        tree = etree.fromstring(blob)
        for p in tree.iter(f"{W}p"):
            yield part, tree, p


def extract_text_units(path: Path) -> list[TextUnit]:
    doc = Document(path)
    units = []
    for i, p in enumerate(_iter_paragraphs(doc)):
        text = paragraph_text(p)
        if text.strip():
            units.append(TextUnit(id=f"p{i}", text=text))
    for i, (part, _tree, p) in enumerate(_extra_parts_paragraphs(path)):
        text = "".join((t.text or "") for t in p.iter(f"{W}t"))
        if text.strip():
            units.append(TextUnit(id=f"extra:{part}:{i}", text=text))
    return units


def scan(path: Path, analyzer, config) -> list:
    findings = []
    for unit in extract_text_units(path):
        findings.extend(detect_unit(analyzer, unit, config))
    return findings


def apply(path: Path, out_path: Path, decisions: dict, analyzer, config, mapping_store) -> None:
    doc = Document(path)
    for p in _iter_paragraphs(doc):
        runs = paragraph_runs(p)
        text = runs_text_and_spans(runs)[0]
        if not text.strip():
            continue
        findings = detect_unit(analyzer, TextUnit(id="tmp", text=text), config)
        apply_findings_to_runs(runs, findings, decisions, mapping_store)
    doc.save(out_path)
    _apply_extra_parts(out_path, analyzer, config, decisions, mapping_store)


def _apply_extra_parts(path: Path, analyzer, config, decisions: dict, mapping_store) -> None:
    with zipfile.ZipFile(path, "r") as zf:
        names = zf.namelist()
        contents = {n: zf.read(n) for n in names}

    any_changed = False
    for part in EXTRA_PARTS:
        if part not in contents:
            continue
        tree = etree.fromstring(contents[part])
        part_changed = False
        for p_elem in tree.iter(f"{W}p"):
            t_elems = list(p_elem.iter(f"{W}t"))
            text = "".join((t.text or "") for t in t_elems)
            if not text.strip():
                continue
            unit = TextUnit(id="tmp", text=text)
            findings = detect_unit(analyzer, unit, config)
            if not findings:
                continue
            runs = [XmlRunAdapter(t) for t in t_elems]
            apply_findings_to_runs(runs, findings, decisions, mapping_store)
            part_changed = True
        if part_changed:
            contents[part] = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
            any_changed = True

    if any_changed:
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for name in names:
                zf.writestr(name, contents[name])
