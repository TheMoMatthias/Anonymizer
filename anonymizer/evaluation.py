"""Measured recall, not asserted recall.

A redaction tool cannot be called "robust" on a feeling. This plants KNOWN names
and identifiers into realistic German bank text and reports how much the
pipeline actually finds -- broken down by the strata that decide the answer.

Why these strata: German NER does not fail at "hard" names, it fails at ORDINARY
ones. Nineteen of the twenty most common German surnames are everyday words
(Müller=miller, Weber=weaver, Bauer=farmer), and the model was trained on
Wikipedia prose, so it leans on sentence context that a form field or a table
cell simply does not have. Measuring one aggregate number would hide exactly
that: foreign surnames in prose score near-perfectly and would mask the
common-noun-in-a-cell case that actually leaks.

Two measurements, deliberately:
  * ISOLATED  - one occurrence, one context, nothing to propagate from. This is
                the pipeline's raw ability to see a name cold. Pessimistic.
  * DOCUMENT  - a realistic letter where the name recurs in several contexts, so
                anchors + document-wide propagation can do their job. This is
                what actually happens to a real file.

HONEST LIMITS -- report these numbers as an UPPER BOUND:
  * The names come from lists we chose, in documents we shaped. Real
    correspondence is messier (OCR noise, typos, nicknames, married names).
  * Recall here is measured against planted PII only; it says nothing about the
    PII we never thought to plant.
  * It is not a labelled sample of the bank's real mail, which is the only
    thing that would settle the question completely.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from .core import detect_unit
from .models import TextUnit

# --- the planted population --------------------------------------------------

# The measured failure mode: German surnames that are also ordinary nouns or
# adjectives. spaCy scores these far worse than exotic ones.
SURNAMES_COMMON_NOUN = [
    "Müller", "Schneider", "Weber", "Bauer", "Klein", "Schwarz", "Richter",
    "Koch", "Braun", "Wolf", "Jung", "Berg", "Fischer", "Vogel", "Hahn",
    "Stein", "Kaiser", "Fuchs", "Sommer", "Winkler",
]
# German surnames that are NOT everyday words -- the control group.
SURNAMES_GERMAN_RARE = [
    "Habermehl", "Bönnighausen", "Rüdenauer", "Osterkamp", "Kretschmar",
    "Wüstefeld", "Nagelschmidt", "Schwanitz",
]
# Non-German surnames, a large share of any German bank's customers.
SURNAMES_FOREIGN = [
    "Öztürk", "Yılmaz", "Nguyen", "Kowalczyk", "Rossi", "Ivanov", "Popescu",
    "Hussein", "Demir", "Petrov",
]
# Nobiliary / patronymic particles. The particle itself identifies nobody, so
# only the CORE is scored (see _identifying_tokens) -- redacting "von [NAME]"
# leaks nothing. What makes these hard is tokenization: the particle is a
# lowercase function word, so a model that requires initial capitals, or a
# shape gate that splits on whitespace, truncates the span mid-name.
SURNAMES_PARTICLE = [
    "von Bergen", "van der Berg", "de Vries", "zu Guttenberg", "von der Leyen",
    "van den Broek", "de la Cruz", "von Hohenstein",
]
# Double-barrelled names. Scored on BOTH halves: catching only "Schmidt" and
# leaving "Rottluff" standing is a disclosure, not a partial success, so the
# harness refuses to credit a half-catch (see _identifying_tokens).
SURNAMES_HYPHENATED = [
    "Schmidt-Rottluff", "Müller-Lüdenscheidt", "Bergmann-Pohl", "Klein-Vogelsang",
    "Sacher-Masoch", "Weizsäcker-Kohl",
]
# Non-ASCII orthography beyond the German umlauts a German model saw in
# training: Vietnamese tone marks, Serbian/Icelandic/Slovak/Romanian letters.
# These break naive normalization and casing rules, and no German gazetteer
# contains them.
SURNAMES_TRANSLITERATED = [
    "Nguyễn", "Đorđević", "Þórsdóttir", "Kováčová", "Ceaușescu", "Åkerlund",
    "Şahin", "García-Ñíguez",
]

# OCR / transcription damage. Scanned correspondence, faxes and legacy exports do
# not contain clean text, and none of the strata above measure what happens when a
# name arrives corrupted. The corruptions here are the ones real OCR actually
# makes: l/1/I confusion, rn->m, ii for an umlaut, O/0, dropped diacritics.
# EXPECTED to score badly on the model and to be carried almost entirely by the
# anchors -- which is the finding, because it says anchors are what protect a
# scanned document.
SURNAMES_OCR_NOISE = [
    "Wlnkler",      # i -> l
    "Mul1er",       # umlaut dropped, l -> 1
    "Miiller",      # ue -> ii, the classic OCR umlaut
    "0sterkamp",    # O -> 0
    "Habermebl",    # h -> b
    "Kretschrnar",  # m -> rn
    "Nagelschmldt", # i -> l
    "Sclnvanitz",   # hw -> lnv
]

SURNAME_STRATA: dict[str, list[str]] = {
    "german_common_noun": SURNAMES_COMMON_NOUN,
    "german_rare": SURNAMES_GERMAN_RARE,
    "foreign": SURNAMES_FOREIGN,
    "particle": SURNAMES_PARTICLE,
    "hyphenated": SURNAMES_HYPHENATED,
    "transliterated": SURNAMES_TRANSLITERATED,
    "ocr_noise": SURNAMES_OCR_NOISE,
}

# Leading particles carry no identifying information on their own.
_PARTICLES = frozenset({"von", "van", "de", "der", "den", "zu", "zur", "zum", "la", "le", "di", "da"})


def _identifying_tokens(surname: str) -> list[str]:
    """The tokens that MUST be redacted for this plant to count as caught.

    Two asymmetric rules, both deliberate:
      * leading particles are DROPPED -- "von Bergen" reduced to "Bergen",
        because leaving a bare "von" behind discloses nothing;
      * hyphen halves are KEPT SEPARATE and all are required -- leaving
        "Rottluff" behind after redacting "Schmidt-" discloses the person.
    Without this split the harness would either punish correct behaviour or
    credit a genuine leak."""
    parts = [p for p in re.split(r"[\s\-]+", surname) if p]
    while len(parts) > 1 and parts[0].lower() in _PARTICLES:
        parts.pop(0)
    return parts

GIVEN_NAMES = ["Björn", "Petra", "Thomas", "Ayşe", "Mehmet", "Anna", "Lukas", "Sofia"]

# Neutral bank prose with no names in it, so a document's only PII is what we
# planted (otherwise propagation could spread a filler name and flatter us).
FILLER = (
    "Die Abrechnung erfolgt quartalsweise gemäß den Allgemeinen Geschäftsbedingungen. "
    "Weitere Unterlagen finden Sie in der Anlage zu diesem Schreiben."
)


# --- the contexts a name appears in ------------------------------------------
# Each returns a line of text containing the surname exactly once.

def _salutation(given: str, surname: str) -> str:
    return f"Sehr geehrter Herr {surname},"


def _prose_full_name(given: str, surname: str) -> str:
    return f"{given} {surname} hat den Vertrag unterzeichnet."


def _prose_oblique(given: str, surname: str) -> str:
    return f"Die Unterlagen wurden von {surname} geprüft und freigegeben."


def _labelled_field(given: str, surname: str) -> str:
    return f"Kunde: {surname}"


def _bare(given: str, surname: str) -> str:
    return surname


def _signature(given: str, surname: str) -> str:
    return f"Mit freundlichen Grüßen {given} {surname}"


# --- the OBLIQUE contexts: a person named with no honorific to anchor on ------
# Every one of these is ordinary German bank correspondence, and every one
# withholds the single cue ("Herr"/"Frau"/"Kunde:") that the anchored patterns
# key on. This is the axis where a redaction tool quietly leaks: the name is
# present, unambiguous to a human, and invisible to a context-gated recognizer.

def _initials(given: str, surname: str) -> str:
    return f"{given[:1]}. {surname} hat den Vorgang gezeichnet."


def _role_reference(given: str, surname: str) -> str:
    return f"Der Einreicher {surname} aus Frankfurt bestätigte den Sachverhalt."


def _distribution_list(given: str, surname: str) -> str:
    # The plant sits BETWEEN two department names, so a model leaning on
    # "items in a list are alike" is actively pushed toward the wrong answer.
    return f"Verteiler: Rechtsabteilung, {surname}, Innenrevision"


def _maiden_name(given: str, surname: str) -> str:
    return f"Die Kundin, geb. {surname}, führt das Konto seit 2011."


def _after_preposition(given: str, surname: str) -> str:
    return f"Nach Rücksprache mit {surname} wurde der Vorgang abgeschlossen."


# --- CASING variants ----------------------------------------------------------
# Forms, mainframe exports and legacy banking systems store names in a single
# case. Every shape gate in this codebase keys off an initial capital, and the
# model was trained on running prose, so a name in one case has none of the
# signal either relies on. Measured: a bare "WINKLER" was detected as NOTHING.

def _bare_upper(given: str, surname: str) -> str:
    return surname.upper()


def _bare_lower(given: str, surname: str) -> str:
    return surname.lower()


def _labelled_upper(given: str, surname: str) -> str:
    return f"KUNDE: {surname.upper()}"


CONTEXTS = {
    "salutation": _salutation,
    "prose_full_name": _prose_full_name,
    "prose_oblique": _prose_oblique,
    "labelled_field": _labelled_field,
    "bare_cell": _bare,
    "signature": _signature,
    "initials": _initials,
    "role_reference": _role_reference,
    "distribution_list": _distribution_list,
    "maiden_name": _maiden_name,
    "after_preposition": _after_preposition,
    "bare_upper": _bare_upper,
    "bare_lower": _bare_lower,
    "labelled_upper": _labelled_upper,
}

# Contexts measured WITHOUT the neutral filler prose.
#
# "bare_cell" has to mean BARE. The harness used to prepend FILLER to every
# probe, including this one -- which handed a WikiNER-trained model two sentences
# of German prose that a lone spreadsheet cell does not have, and made the
# hardest and most common shape in the user's real workbooks score like prose.
# Every bare_cell figure reported before this was optimistic.
_NO_FILLER_CONTEXTS = frozenset({"bare_cell", "bare_upper", "bare_lower"})


def probe_text(context: str, given: str, surname: str) -> str:
    """The exact text one probe feeds the pipeline. Separate from the measuring
    loop so the "is the bare case actually bare" property is directly testable."""
    line = CONTEXTS[context](given, surname)
    return line if context in _NO_FILLER_CONTEXTS else f"{FILLER} {line}"

# Structured identifiers: these have checksums or hard patterns, so they should
# score near 1.0. A dip here is a much louder alarm than a dip on names.
STRUCTURED_PROBES: dict[str, tuple[str, str]] = {
    "IBAN": ("DE89370400440532013000", "Bitte überweisen Sie auf IBAN {v} zugunsten des Kontos."),
    "STEUER_ID": ("86095742719", "Die Steuer-ID lautet {v} laut Bescheid."),
    "EMAIL": ("b.mueller@example.de", "Antworten Sie bitte an {v} zurück."),
    "PHONE_DE": ("0170 1234567", "Telefon: {v} für Rückfragen."),
    "ADDRESS": ("Königsallee 3", "Anschrift: {v} in der Akte."),
    "PLZ_CITY": ("50667 Köln", "Wohnort ist {v} laut Unterlagen."),
    "BIC": ("COBADEFFXXX", "Zahlung an BIC: {v} veranlassen."),
    "DATE_DOB": ("15.03.1980", "Geburtsdatum: {v} des Kunden."),
    "SV_NUMMER": ("65170839J003", "Die Versicherungsnummer {v} ist hinterlegt."),
    # HELD-OUT Art. 9 values. These deliberately do NOT appear in the recognizers'
    # word lists: probing with "evangelisch" / "schwerbehindert" -- both literal
    # entries in the very alternation this measures -- only proved the list
    # contains its own contents, which is tuning the recognizer to the benchmark.
    # A held-out value can only be found through the generic LABEL:VALUE
    # mechanism, so the number means something. The BARE row is expected to miss
    # them, and that miss is the honest finding: an unlabelled free-text health or
    # religion value that is not on the list is not detected.
    "ART9_RELIGION": ("neuapostolisch", "Konfession: {v} laut Stammdaten."),
    "ART9_HEALTH": ("Bandscheibenvorfall", "Diagnose: {v} laut Attest."),
}


# --- Art. 9 stated OBLIQUELY --------------------------------------------------
# GDPR Art. 9 data carries the heaviest consequence of any category here, and the
# recognizers find it two ways: a literal word list, or a LABEL:VALUE anchor
# ("Konfession: …"). Real correspondence supplies NEITHER. It states the fact in
# a plain sentence, and the sensitive part is a perfectly ordinary noun.
#
# The needle is the single token that MUST be removed for the fact to stay
# private. Leave "Rollstuhl" standing and the reader learns the customer's
# disability regardless of what else was redacted.
#
# These are EXPECTED to score badly. That is the finding, not a harness bug: it
# says the word-list approach does not generalize, and it quantifies by how much.
ART9_OBLIQUE_PROBES: dict[str, tuple[str, str]] = {
    "health_disability": ("Rollstuhl", "Der Kunde ist seit dem Unfall dauerhaft auf den Rollstuhl angewiesen."),
    "health_treatment": ("Chemotherapie", "Die Termine mussten wegen der laufenden Chemotherapie verschoben werden."),
    "health_psych": ("Burnout", "Er ist nach einem Burnout seit März nicht im Dienst."),
    "religion_practice": ("Moschee", "Er besucht jeden Freitag die Moschee in der Innenstadt."),
    "religion_observance": ("Ramadan", "Während des Ramadan bittet sie um spätere Gesprächstermine."),
    "union_membership": ("Betriebsrat", "Er wurde in den Betriebsrat gewählt und ist dafür freigestellt."),
    "union_strike": ("Streikgeld", "Für die Ausfalltage wurde Streikgeld an das Mitglied ausgezahlt."),
    "sex_life_partner": ("Ehefrau", "Sie lebt mit ihrer Ehefrau in einer eingetragenen Partnerschaft."),
    "ethnic_origin": ("Kontingentflüchtlinge", "Die Familie kam 1994 als Kontingentflüchtlinge nach Deutschland."),
    "ethnic_language": ("Romanes", "Das Beratungsgespräch wurde auf Romanes geführt."),
    "political_party": ("Grünen", "Er kandidierte bei der Kommunalwahl für die Grünen."),
    "biometric": ("Fingerabdruck", "Die Freigabe erfolgt per Fingerabdruck des Kontoinhabers."),
}

# HELD-OUT Art. 9 probes: every value below is deliberately ABSENT from every
# shipped word list, so it can only be caught by a disclosure FRAME. Kept as its
# own stratum because it answers a different question from the probes above --
# not "does coverage improve" but "does the mechanism GENERALIZE", which is the
# only thing that distinguishes a real improvement from tuning the recognizer to
# its own benchmark. If this stratum scores while the lists stay unchanged, the
# frames are doing real work.
ART9_HELDOUT_PROBES: dict[str, tuple[str, str]] = {
    "frame_leidet_an": ("Sarkoidose", "Der Kunde leidet an einer chronischen Sarkoidose."),
    "frame_erkrankt_an": ("Tuberkulose", "Die Mitarbeiterin ist im Mai an Tuberkulose erkrankt."),
    "frame_angewiesen": ("Beatmungsgerät", "Sie ist nachts auf das Beatmungsgerät angewiesen."),
    "frame_krank_wegen": ("Kniearthroskopie", "Er ist krankgeschrieben wegen einer Kniearthroskopie."),
    "frame_konvertiert": ("Buddhismus", "Er konvertierte im Jahr 2018 zum Buddhismus."),
    "frame_gewaehlt_in": ("Wirtschaftsausschuss", "Sie wurde in den Wirtschaftsausschuss gewählt."),
    "frame_mitglied_im": ("Sprecherausschuss", "Er ist Mitglied im Sprecherausschuss der leitenden Angestellten."),
    "frame_staemmig": ("kasachischstämmig", "Die Familie ist kasachischstämmig und seit 1998 hier."),
}


# --- structured-cell traps ----------------------------------------------------
# The shapes a real "database" workbook stores names in, where the column HEADER
# gives detection nothing to work with. Each trap is a column; each column gets
# its OWN disjoint set of surnames so a hit attributes unambiguously back to one
# trap (a ScanResult groups by value, so shared names would blur the columns).
_TRAP_COLUMNS: dict[str, tuple[str, str]] = {
    # trap name -> (header text, cell template with {n} for the name)
    "lying_header": ("Status", "{n}"),
    "opaque_header": ("Feld_7", "{n}"),
    "no_header": ("", "{n}"),
    "multi_value_cell": ("Vermerk", "{n}; intern geprüft"),
    "id_shaped_cell": ("Referenz", "K-{n}-2024"),
    "initials_cell": ("Kürzel", "B. {n}"),
}


def _trap_partition() -> dict[str, list[str]]:
    """Deal every planted surname out across the traps, round-robin over the
    strata so no trap is accidentally easier than another (all-foreign would
    flatter, all-common-noun would damn). Deterministic -- the harness must
    produce the same workbook on every run for its numbers to be comparable."""
    interleaved: list[str] = []
    pools = [list(v) for v in SURNAME_STRATA.values()]
    for i in range(max(len(p) for p in pools)):
        for pool in pools:
            if i < len(pool):
                interleaved.append(pool[i])
    traps = list(_TRAP_COLUMNS)
    out: dict[str, list[str]] = {t: [] for t in traps}
    for i, name in enumerate(interleaved):
        out[traps[i % len(traps)]].append(name)
    return out


@dataclass
class StratumResult:
    stratum: str
    context: str
    found: int = 0
    total: int = 0
    missed: list[str] = field(default_factory=list)

    @property
    def recall(self) -> float:
        return self.found / self.total if self.total else 0.0


def _whole_token(needle: str, haystack: str) -> bool:
    """Whole-token match, NOT substring: a common-noun surname ("Berg", "Koch")
    must not count as found just because it is a substring of an unrelated finding
    ("Bergstraße", a "…berg" ORG) -- that over-reports recall for exactly the
    stratum this harness exists to measure honestly.

    CASE-INSENSITIVE: redacting "WINKLER" removes the plant just as completely as
    redacting "Winkler", so scoring the casing strata case-sensitively would
    report a correct catch as a leak."""
    return re.search(rf"(?<!\w){re.escape(needle)}(?!\w)", haystack, re.IGNORECASE) is not None


def _found(findings, needle: str) -> bool:
    """A plant counts as found only when EVERY identifying token is covered by
    the union of the findings. The entity TYPE is deliberately not checked: for
    redaction, catching "Bauer" as NER_MISC rather than PERSON still removes it
    -- the leak is what matters.

    The union (rather than "some single finding contains the whole string") is
    what makes multi-token plants scoreable at all: a hyphenated name legitimately
    arrives as two adjacent spans, and that is a full catch, not a miss."""
    tokens = _identifying_tokens(needle)
    return all(any(_whole_token(tok, f.value) for f in findings) for tok in tokens)


def measure_isolated(analyzer, config: dict) -> list[StratumResult]:
    """One name, one context, no other occurrence to propagate from -- the
    pipeline's cold-read ability. Runs on text directly (no file I/O), which is
    valid here precisely because propagation has nothing to work with."""
    cfg = {**config, "languages": ["de"]}
    results: list[StratumResult] = []
    for stratum, surnames in SURNAME_STRATA.items():
        for ctx_name in CONTEXTS:
            r = StratumResult(stratum=stratum, context=ctx_name)
            for i, surname in enumerate(surnames):
                given = GIVEN_NAMES[i % len(GIVEN_NAMES)]
                text = probe_text(ctx_name, given, surname)
                findings = detect_unit(analyzer, TextUnit("u1", text), cfg)
                r.total += 1
                if _found(findings, surname):
                    r.found += 1
                else:
                    r.missed.append(surname)
            results.append(r)
    return results


def measure_structured(analyzer, config: dict) -> list[StratumResult]:
    """Every identifier is measured TWICE: once inside a sentence that carries
    the recognizer's context word, and once BARE -- alone in a cell, no label
    anywhere, which is the shape a "database" workbook actually stores. The
    in-context row flatters every context-gated recognizer; the bare row is the
    one that predicts what leaks out of the user's real files."""
    cfg = {**config, "languages": ["de"]}
    results: list[StratumResult] = []
    for stratum, bare in (("structured", False), ("structured_bare", True)):
        for label, (value, template) in STRUCTURED_PROBES.items():
            r = StratumResult(stratum=stratum, context=label, total=1)
            text = value if bare else template.format(v=value)
            findings = detect_unit(analyzer, TextUnit("u1", text), cfg)
            # Compare space-insensitively: a recognizer may claim a reformatted span.
            flat = {f.value.replace(" ", "") for f in findings}
            if any(value.replace(" ", "") in f for f in flat):
                r.found = 1
            else:
                r.missed.append(value)
            results.append(r)
    return results


def _actionable_values(analyzer, cfg: dict, text: str) -> list[str]:
    """What a REVIEWER would actually be shown for a piece of text.

    Deliberately routed through build_scan_result rather than reading detect_unit
    directly: corroboration-only DEMOTES uncorroborated guesses out of the
    actionable list, and a measurement that skips that step reports coverage the
    tool does not really give. This matters most for Art. 9 -- several of those
    probes were only ever "caught" because spaCy mistyped a capitalized German
    noun as a PERSON, and such a hit is exactly what demotion removes."""
    from .core import build_scan_result

    unit = TextUnit("u1", text)
    findings = detect_unit(analyzer, unit, cfg)
    result = build_scan_result(findings, [unit], cfg)
    return [g.value for g in result.all_actionable()]


class _ValueFinding:
    """Minimal stand-in so _found() can score plain strings."""

    __slots__ = ("value",)

    def __init__(self, value: str) -> None:
        self.value = value


def _measure_art9_probes(analyzer, config: dict, probes: dict, stratum: str) -> list[StratumResult]:
    cfg = {**config, "languages": ["de"]}
    results: list[StratumResult] = []
    for label, (needle, text) in probes.items():
        r = StratumResult(stratum=stratum, context=label, total=1)
        claimed = [_ValueFinding(v) for v in _actionable_values(analyzer, cfg, f"{FILLER} {text}")]
        if _found(claimed, needle):
            r.found = 1
        else:
            r.missed.append(needle)
        results.append(r)
    return results


def measure_art9_oblique(analyzer, config: dict) -> list[StratumResult]:
    """Art. 9 facts stated in plain sentences, with no label and no list word.
    Scored on the one token that has to disappear for the fact to stay private."""
    return _measure_art9_probes(analyzer, config, ART9_OBLIQUE_PROBES, "art9_oblique")


def measure_art9_heldout(analyzer, config: dict) -> list[StratumResult]:
    """Art. 9 values that are in NO shipped word list, so only a frame can catch
    them. This is the generalization check -- see ART9_HELDOUT_PROBES."""
    return _measure_art9_probes(analyzer, config, ART9_HELDOUT_PROBES, "art9_heldout")


# Art. 9 values as a spreadsheet stores them: a bare cell, under a header that
# gives detection nothing. This is the intersection of the tool's two weakest
# areas -- special-category data and structured cells -- and nothing measured it.
# The values are HELD OUT of the shipped word lists on purpose where possible, so
# a hit means the mechanism generalizes rather than that the list contains itself.
ART9_CELL_PROBES: dict[str, list[str]] = {
    "health": ["Dialysepatient", "Rollstuhlfahrer", "Bandscheibenvorfall", "schwerbehindert"],
    "religion": ["neuapostolisch", "alevitisch", "russisch-orthodox", "Zeuge Jehovas"],
    "union": ["ver.di-Mitglied", "IG Metall", "Betriebsratsvorsitzender", "Streikkasse"],
    "ethnic": ["Sinti", "kurdischer Herkunft", "Kontingentflüchtling", "Aramäer"],
}


def measure_art9_structured(analyzer, config: dict, workdir: Path) -> list[StratumResult]:
    """Art. 9 values in bare spreadsheet cells, under headers that say nothing."""
    import openpyxl

    from .pipeline import scan_document

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Stammdaten"
    ws.cell(row=1, column=1, value="Vorgang_ID")
    headers = ["Merkmal_1", "Status", "Feld_3", "Bemerkung"]
    for col, (cat, header) in enumerate(zip(ART9_CELL_PROBES, headers), start=2):
        ws.cell(row=1, column=col, value=header)
        for row, value in enumerate(ART9_CELL_PROBES[cat], start=2):
            ws.cell(row=row, column=col, value=value)
    for row in range(2, 6):
        ws.cell(row=row, column=1, value=f"VG-{2000 + row}")
    path = workdir / "art9_cells.xlsx"
    wb.save(path)
    wb.close()

    claimed = [g.value for g in scan_document(path, analyzer, config).all_actionable()]
    results: list[StratumResult] = []
    for cat, values in ART9_CELL_PROBES.items():
        r = StratumResult(stratum="art9_structured", context=cat)
        for value in values:
            r.total += 1
            if any(_whole_token(tok, v) for v in claimed for tok in _identifying_tokens(value)):
                r.found += 1
            else:
                r.missed.append(value)
        results.append(r)
    return results


# A name wrapped inside an identifier, a path or an address. The name is fully
# legible to any reader and there is nothing in this codebase that looks INSIDE a
# delimited token for one.
EMBEDDED_TEMPLATES: dict[str, str] = {
    "id_hyphen": "Referenz: K-{n}-2024",
    "id_underscore": "Ablage: AKTE_{n}_2024",
    "ticket_ref": "Vorgang TICKET-4711-{n} ist offen.",
    "unc_path": r"Pfad: \\fileserver\Kunden\{n}\2024\Vertrag.pdf",
    "filename": "Datei: Vertrag_{n}_final_v2.pdf",
}


def measure_embedded_identifiers(analyzer, config: dict, workdir: Path) -> list[StratumResult]:
    """Names inside identifiers, paths and filenames, in a document that ALSO
    names the person normally.

    The anchor sentence is not a kindness -- it is what makes the measurement mean
    anything. An identifier probed in total isolation ("K-Winkler-2024" and
    nothing else) is not solvable by any mechanism this tool has or could
    reasonably have: there is no signal separating that from a product code, and
    a surname list broad enough to decide it would flag every capitalized token in
    every id. What IS solvable, and what a real document always provides, is the
    same name appearing normally elsewhere -- so the question this asks is whether
    a KNOWN name is still found once it is wrapped in an identifier. Scored on the
    embedded occurrence only, by requiring a count of 2."""
    from docx import Document

    from .pipeline import scan_document

    partition = _trap_partition()
    pool = [n for names in partition.values() for n in names]
    results: list[StratumResult] = []
    for label, template in EMBEDDED_TEMPLATES.items():
        r = StratumResult(stratum="embedded", context=label)
        for i, surname in enumerate(pool[:12]):
            doc = Document()
            doc.add_paragraph(_salutation("", surname))  # establishes the name
            doc.add_paragraph(FILLER)
            doc.add_paragraph(template.format(n=surname))
            path = workdir / f"embedded_{label}_{i}.docx"
            doc.save(path)
            r.total += 1
            # 2 occurrences planted: the salutation and the identifier. Anything
            # less than both means the embedded one was not reached.
            if _occurrences_caught(scan_document(path, analyzer, config), surname) >= 2:
                r.found += 1
            else:
                r.missed.append(surname)
        results.append(r)
    return results


def measure_workbook_traps(analyzer, config: dict, workdir: Path) -> list[StratumResult]:
    """Names in spreadsheet cells whose column header offers no help at all.

    This is the shape the user's real files are full of and the one the
    letter-shaped measurements above cannot see: there is no salutation to anchor
    on, no surrounding sentence for a WikiNER-trained model to lean on, and no
    people-word in the header for the whole-cell override to fire on. Scored per
    planted cell."""
    import openpyxl

    from .pipeline import scan_document

    partition = _trap_partition()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Vorgaenge"
    traps = list(_TRAP_COLUMNS)
    for col, trap in enumerate(traps, start=1):
        header, template = _TRAP_COLUMNS[trap]
        if header:
            ws.cell(row=1, column=col, value=header)
        for row, name in enumerate(partition[trap], start=2):
            ws.cell(row=row, column=col, value=template.format(n=name))
    # A neutral first column so the sheet reads as a real table rather than a
    # column of bare names -- the latter is an easier problem than reality.
    ws.insert_cols(1)
    ws.cell(row=1, column=1, value="Vorgang_ID")
    for row in range(2, ws.max_row + 1):
        ws.cell(row=row, column=1, value=f"VG-{1000 + row}")
    path = workdir / "trap_workbook.xlsx"
    wb.save(path)
    wb.close()

    result = scan_document(path, analyzer, config)
    claimed = [g.value for g in result.all_actionable()]
    results: list[StratumResult] = []
    for trap in traps:
        r = StratumResult(stratum="structured_trap", context=trap)
        for name in partition[trap]:
            r.total += 1
            tokens = _identifying_tokens(name)
            if all(any(_whole_token(tok, v) for v in claimed) for tok in tokens):
                r.found += 1
            else:
                r.missed.append(name)
        results.append(r)
    return results


def measure_documents(analyzer, config: dict, workdir: Path) -> list[StratumResult]:
    """A realistic letter: the name recurs across salutation, prose, a labelled
    field and a bare cell. Every occurrence must be caught -- this is where the
    anchors seed a name and propagation spreads it to the units NER cannot see.
    Scored per OCCURRENCE, not per document, so a partial catch cannot pass."""
    from docx import Document

    from .pipeline import scan_document

    results: list[StratumResult] = []
    for stratum, surnames in SURNAME_STRATA.items():
        r = StratumResult(stratum=stratum, context="full_letter_occurrences")
        for i, surname in enumerate(surnames):
            given = GIVEN_NAMES[i % len(GIVEN_NAMES)]
            doc = Document()
            doc.add_paragraph(_salutation(given, surname))
            doc.add_paragraph(FILLER)
            doc.add_paragraph(_prose_oblique(given, surname))
            doc.add_paragraph(_labelled_field(given, surname))
            table = doc.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = _bare(given, surname)
            doc.add_paragraph(_signature(given, surname))
            path = workdir / f"letter_{stratum}_{i}.docx"
            doc.save(path)

            planted = 5  # salutation, oblique, labelled, bare cell, signature
            result = scan_document(path, analyzer, config)
            caught = _occurrences_caught(result, surname)
            r.total += planted
            r.found += min(caught, planted)
            if caught < planted:
                r.missed.append(f"{surname}({caught}/{planted})")
        results.append(r)
    return results


def _occurrences_caught(result, surname: str) -> int:
    """How many occurrences of a plant were claimed, scored per identifying token.

    The count is the MINIMUM across the tokens, not the sum: for "Schmidt-Rottluff"
    a run that claimed every "Schmidt" and no "Rottluff" has redacted nothing
    safely, so it scores 0 rather than half. Summing would let a systematic
    half-catch report as partial success."""
    tokens = _identifying_tokens(surname)
    groups = list(result.all_actionable())
    per_token = [
        sum(g.count for g in groups if _whole_token(tok, g.value)) for tok in tokens
    ]
    return min(per_token) if per_token else 0


def measure_unanchored_documents(analyzer, config: dict, workdir: Path) -> list[StratumResult]:
    """The hardest realistic document: an internal memo that names a person
    several times and NEVER once with an honorific or a "Kunde:" label.

    Why this is the measurement that matters most: the anchored patterns are what
    seed document-wide propagation, so the flattering 98% in the section above is
    really "one salutation rescued every other occurrence". Remove the salutation
    -- as any internal note, meeting minute or ticket comment does -- and
    propagation has nothing to spread. This section is the tool's true floor."""
    from docx import Document

    from .pipeline import scan_document

    contexts = ["after_preposition", "role_reference", "initials", "distribution_list", "bare_cell"]
    results: list[StratumResult] = []
    for stratum, surnames in SURNAME_STRATA.items():
        r = StratumResult(stratum=stratum, context="unanchored_occurrences")
        for i, surname in enumerate(surnames):
            given = GIVEN_NAMES[i % len(GIVEN_NAMES)]
            doc = Document()
            doc.add_paragraph(FILLER)
            for ctx in contexts[:-1]:
                doc.add_paragraph(CONTEXTS[ctx](given, surname))
            table = doc.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = _bare(given, surname)
            path = workdir / f"memo_{stratum}_{i}.docx"
            doc.save(path)

            planted = len(contexts)
            caught = _occurrences_caught(scan_document(path, analyzer, config), surname)
            r.total += planted
            r.found += min(caught, planted)
            if caught < planted:
                r.missed.append(f"{surname}({caught}/{planted})")
        results.append(r)
    return results


def format_report(sections: dict[str, list[StratumResult]]) -> str:
    lines = ["", "=" * 74, "RECALL REPORT  (planted PII -- treat as an UPPER BOUND)", "=" * 74]
    for title, results in sections.items():
        lines.append("")
        lines.append(f"--- {title} ---")
        lines.append(f"{'stratum':<22}{'context':<26}{'recall':>9}  {'n':>5}")
        for r in results:
            flag = "" if r.recall >= 0.9 else ("  <-- WEAK" if r.recall >= 0.5 else "  <-- LEAKING")
            lines.append(f"{r.stratum:<22}{r.context:<26}{r.recall:>8.0%}  {r.total:>5}{flag}")
        total = sum(r.total for r in results)
        found = sum(r.found for r in results)
        lines.append(f"{'':<22}{'OVERALL':<26}{(found / total if total else 0):>8.0%}  {total:>5}")
    misses = [
        f"  {r.stratum}/{r.context}: {', '.join(r.missed)}"
        for results in sections.values()
        for r in results
        if r.missed
    ]
    if misses:
        lines += ["", "MISSED (what still leaks):"] + misses
    lines.append("")
    return "\n".join(lines)
